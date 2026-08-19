#!/usr/bin/env python3
"""Generate Forus-branded competency & expectations map (Word) for sales managers."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Карта_компетенций_менеджер_продуктовый_запуск.docx"
OUT_ROOT = ROOT / "Карта_компетенций_менеджер_продуктовый_запуск.docx"
ART_RU = Path("/opt/cursor/artifacts/Карта_компетенций_менеджер_продуктовый_запуск.docx")
ART_EN = Path("/opt/cursor/artifacts/Karta_kompetenciy_menedzher_produktovyy_zapusk.docx")

# Brand: ГК Форус (светлый корпоративный стиль)
YELLOW = RGBColor(0xFC, 0xCD, 0x68)
NEAR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x76, 0x76, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT = RGBColor(0xFF, 0xF8, 0xE8)
HEADER_BG = YELLOW
REALITY_BG = RGBColor(0xFA, 0xFA, 0xFA)
LABEL_BG = RGBColor(0xFF, 0xF3, 0xD6)
COMP_BG = RGBColor(0xF0, 0xF7, 0xFB)  # лёгкий голубой для столбца компетенций
FONT = "Verdana"

CORP_COMPETENCIES = [
    "Исполнительность",
    "Клиентоориентированность",
    "Устная коммуникация",
    "Письменная коммуникация",
    "Коммуникабельность",
    "Стрессоустойчивость",
    "Командность",
    "Самоорганизация",
    "Ответственность",
    "Ориентация на результат",
    "Адаптивность",
    "Настойчивость",
    "Решительность",
    "Системное мышление",
    "Проактивность / инновационность / инициатива",
    "Внимание к деталям",
    "Обучаемость",
    "Пунктуальность",
]

# (действие, список корпоративных компетенций)
ACTIONS: list[tuple[str, list[str]]] = [
    (
        "Чёткое соблюдение скрипта в холодной базе (структура звонка, ключевые вопросы, фиксация результата)",
        ["Исполнительность", "Устная коммуникация", "Внимание к деталям", "Самоорганизация"],
    ),
    (
        "Чёткое соблюдение скрипта в тёплой базе (развитие интереса, квалификация, следующий шаг)",
        ["Исполнительность", "Клиентоориентированность", "Устная коммуникация", "Внимание к деталям"],
    ),
    (
        "Чёткое соблюдение скрипта в горячей базе (закрытие сделки, работа с решением, дожим договорённостей)",
        ["Исполнительность", "Ориентация на результат", "Настойчивость", "Решительность", "Устная коммуникация"],
    ),
    (
        "Быстрое установление контакта и поддержание диалога с незнакомым ЛПР в холодной / тёплой базе",
        ["Коммуникабельность", "Устная коммуникация", "Клиентоориентированность", "Стрессоустойчивость"],
    ),
    (
        "Быстрое изучение нового продукта: ценность, отличия, ограничения, типовые сценарии применения",
        ["Обучаемость", "Адаптивность", "Внимание к деталям", "Системное мышление"],
    ),
    (
        "Знание особенностей продукта и умение объяснить их простым языком под задачу клиента",
        ["Устная коммуникация", "Клиентоориентированность", "Обучаемость", "Внимание к деталям"],
    ),
    (
        "Проработка гипотез при выводе продукта на рынок (тестирование оффера, сбор обратной связи)",
        ["Системное мышление", "Проактивность / инновационность / инициатива", "Ориентация на результат", "Обучаемость"],
    ),
    (
        "Выявление потребностей клиента и перевод их в продуктовую ценность",
        ["Клиентоориентированность", "Устная коммуникация", "Системное мышление", "Внимание к деталям"],
    ),
    (
        "Презентация продукта: структура, аргументы, акцент на выгоде, а не на функциях",
        ["Устная коммуникация", "Клиентоориентированность", "Системное мышление", "Ориентация на результат"],
    ),
    (
        "Отработка возражений по скрипту и с опорой на продуктовую экспертизу",
        ["Стрессоустойчивость", "Устная коммуникация", "Настойчивость", "Решительность", "Обучаемость"],
    ),
    (
        "Квалификация лида: понимание готовности, роли ЛПР, сроков и бюджета",
        ["Системное мышление", "Внимание к деталям", "Ориентация на результат", "Решительность"],
    ),
    (
        "Работа с CRM: корректная и своевременная фиксация статусов, комментариев и следующих шагов",
        ["Письменная коммуникация", "Внимание к деталям", "Исполнительность", "Пунктуальность", "Ответственность"],
    ),
    (
        "Подготовка писем, КП, саммари звонка и статусов по запуску продукта в понятной письменной форме",
        ["Письменная коммуникация", "Внимание к деталям", "Клиентоориентированность", "Исполнительность"],
    ),
    (
        "Обратная связь продуктовой команде / тимлиду по рынку, возражениям и гипотезам",
        ["Письменная коммуникация", "Командность", "Проактивность / инновационность / инициатива", "Системное мышление"],
    ),
    (
        "Адаптация при смене продукта, оффера или скрипта без потери качества звонка",
        ["Адаптивность", "Обучаемость", "Стрессоустойчивость", "Исполнительность"],
    ),
    (
        "Управление воронкой: приоритезация базы, доведение контактов до результата",
        ["Самоорганизация", "Ориентация на результат", "Системное мышление", "Ответственность", "Настойчивость"],
    ),
    (
        "Соблюдение слотов демо, дедлайнов по задачам запуска и договорённых сроков обратного контакта",
        ["Пунктуальность", "Ответственность", "Исполнительность", "Клиентоориентированность"],
    ),
    (
        "Самостоятельное планирование дня: объём касаний, приоритеты по базе и подготовка к новым продуктам",
        ["Самоорганизация", "Ответственность", "Ориентация на результат", "Пунктуальность"],
    ),
    (
        "Принятие решения о следующем шаге в сделке (демо / КП / эскалация / отказ) без затягивания",
        ["Решительность", "Ориентация на результат", "Системное мышление", "Ответственность"],
    ),
    (
        "Дожим договорённостей и повторные касания по тёплой / горячей базе до ясности статуса",
        ["Настойчивость", "Ориентация на результат", "Стрессоустойчивость", "Ответственность"],
    ),
    (
        "Обмен лучшими практиками с коллегами при запуске продукта и поддержка команды",
        ["Командность", "Проактивность / инновационность / инициатива", "Коммуникабельность", "Обучаемость"],
    ),
    (
        "Предложения по улучшению скрипта, оффера и гипотез на основе практики звонков",
        ["Проактивность / инновационность / инициатива", "Системное мышление", "Командность", "Обучаемость"],
    ),
    (
        "Сохранение рабочего тона и стандартов качества при отказах и высокой нагрузке холодной базы",
        ["Стрессоустойчивость", "Клиентоориентированность", "Исполнительность", "Коммуникабельность"],
    ),
]

EXPECTATIONS: list[tuple[str, list[str]]] = [
    (
        "Самостоятельность в отработке возражений без эскалации типовых ситуаций руководителю",
        ["Решительность", "Стрессоустойчивость", "Обучаемость", "Ответственность"],
    ),
    (
        "Дисциплина по скрипту: без самовольных отступлений, искажения смысла и «своей версии»",
        ["Исполнительность", "Внимание к деталям", "Ответственность"],
    ),
    (
        "Выход на уверенную работу с новым продуктом в согласованные сроки после обучения",
        ["Обучаемость", "Адаптивность", "Пунктуальность", "Ориентация на результат"],
    ),
    (
        "Выполнение норм активности (звонки / касания) при сохранении качества диалога",
        ["Исполнительность", "Самоорганизация", "Ориентация на результат", "Стрессоустойчивость"],
    ),
    (
        "Качественное ведение CRM: полнота, достоверность, понятные следующие действия",
        ["Письменная коммуникация", "Внимание к деталям", "Ответственность", "Пунктуальность"],
    ),
    (
        "Инициативность: предложения по улучшению скрипта, оффера и гипотез на основе практики",
        ["Проактивность / инновационность / инициатива", "Системное мышление", "Командность"],
    ),
    (
        "Стабильная конверсия по этапам воронки в рамках целевых показателей отдела",
        ["Ориентация на результат", "Настойчивость", "Системное мышление", "Ответственность"],
    ),
    (
        "Устойчивость в работе с холодной базой: сохранение энергии, тона и стандартов",
        ["Стрессоустойчивость", "Коммуникабельность", "Исполнительность", "Клиентоориентированность"],
    ),
    (
        "Своевременная эскалация нетиповых кейсов и рисков по сделке / продукту",
        ["Ответственность", "Решительность", "Командность", "Системное мышление"],
    ),
    (
        "Командное взаимодействие: обмен лучшими практиками и поддержка коллег при запуске",
        ["Командность", "Коммуникабельность", "Проактивность / инновационность / инициатива"],
    ),
    (
        "Клиентоориентированность: уважительный тон, честность по возможностям продукта",
        ["Клиентоориентированность", "Устная коммуникация", "Ответственность"],
    ),
    (
        "Ответственность за результат: доведение договорённостей до конкретного следующего шага",
        ["Ответственность", "Ориентация на результат", "Настойчивость", "Пунктуальность"],
    ),
    (
        "Соблюдение времени демо, дедлайнов по задачам запуска и обещанных клиенту сроков ответа",
        ["Пунктуальность", "Клиентоориентированность", "Исполнительность", "Ответственность"],
    ),
    (
        "Понятная письменная коммуникация: письма, КП, саммари и комментарии без потери смысла",
        ["Письменная коммуникация", "Внимание к деталям", "Клиентоориентированность"],
    ),
]


def set_run_font(run, size=10, bold=False, color=NEAR_BLACK, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), font)


def set_cell_shading(cell, color: RGBColor):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(shd)


def set_cell_borders(cell, color="D0D0D0", sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)
    tc_pr.append(borders)


def set_cell_margins(cell, top=50, bottom=50, left=60, right=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for name, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    for old in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(old)
    tc_pr.append(tc_mar)


def set_vertical_align(cell, val="center"):
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), val)
    for old in tc_pr.findall(qn("w:vAlign")):
        tc_pr.remove(old)
    tc_pr.append(v_align)


def clear_paragraph(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1


def write_cell(
    cell,
    text,
    size=9,
    bold=False,
    color=NEAR_BLACK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    fill=None,
):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if fill is not None:
        set_cell_shading(cell, fill)
    set_cell_borders(cell)
    set_cell_margins(cell)
    set_vertical_align(cell, "center")


def add_horizontal_line(paragraph, color="FCCD68", size="24"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_table_full_width(table):
    tbl_pr = table._tbl.tblPr
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    for old in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(old)
    tbl_pr.append(tbl_w)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def format_comps(comps: list[str]) -> str:
    return " · ".join(comps)


def covered_competencies() -> set[str]:
    covered: set[str] = set()
    for _, comps in ACTIONS + EXPECTATIONS:
        covered.update(comps)
    return covered


def build() -> Path:
    missing = [c for c in CORP_COMPETENCIES if c not in covered_competencies()]
    if missing:
        raise RuntimeError(f"Competencies not covered by actions: {missing}")

    doc = Document()
    section = doc.sections[0]
    # Landscape A4 — удобнее для 3 столбцов
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    brand = doc.add_table(rows=1, cols=2)
    set_table_full_width(brand)
    write_cell(brand.rows[0].cells[0], "ФОРУС", size=16, bold=True, fill=YELLOW)
    write_cell(
        brand.rows[0].cells[1],
        "Группа компаний  ·  Отдел продуктового запуска",
        size=9,
        fill=YELLOW,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    set_col_widths(brand, [6.0, 21.0])
    for cell in brand.rows[0].cells:
        set_cell_borders(cell, color="E8B43A", sz="4")
        set_cell_margins(cell, top=70, bottom=70, left=90, right=90)

    title = doc.add_paragraph()
    clear_paragraph(title)
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(3)
    set_run_font(title.add_run("Карта компетенций и ожиданий сотрудника"), size=14, bold=True)
    add_horizontal_line(title, color="FCCD68", size="28")

    subtitle = doc.add_paragraph()
    clear_paragraph(subtitle)
    subtitle.paragraph_format.space_before = Pt(4)
    subtitle.paragraph_format.space_after = Pt(8)
    set_run_font(
        subtitle.add_run("Менеджер по продажам  ·  Отдел продуктового запуска"),
        size=10,
        color=GRAY,
    )

    meta = doc.add_table(rows=3, cols=4)
    set_table_full_width(meta)
    meta_rows = [
        [("Сотрудник", True), ("", False), ("Дата оценки", True), ("", False)],
        [("Должность", True), ("Менеджер по продажам", False), ("Оценщик", True), ("", False)],
        [("Отдел", True), ("Продуктовый запуск", False), ("Период", True), ("", False)],
    ]
    for i, row_data in enumerate(meta_rows):
        for j, (text, is_label) in enumerate(row_data):
            cell = meta.rows[i].cells[j]
            if is_label:
                write_cell(cell, text, size=8, bold=True, fill=LABEL_BG)
            else:
                display = text if text else "________________________________"
                write_cell(
                    cell,
                    display,
                    size=9,
                    color=NEAR_BLACK if text else GRAY,
                    fill=WHITE,
                )
    set_col_widths(meta, [3.5, 10.0, 3.5, 10.0])

    note = doc.add_paragraph()
    clear_paragraph(note)
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(4)
    set_run_font(
        note.add_run(
            "Как читать таблицу: «Действие» — наблюдаемое поведение в отделе продуктового запуска; "
            "«Корпоративные компетенции» — к каким компетенциям Форус относится действие (может быть несколько); "
            "«Как у сотрудника в реальности» — факт / примеры / оценка (шкала 1–4: не соответствует → превосходит)."
        ),
        size=8,
        color=GRAY,
    )

    # Reference list of corporate competencies
    ref = doc.add_paragraph()
    clear_paragraph(ref)
    ref.paragraph_format.space_before = Pt(2)
    ref.paragraph_format.space_after = Pt(8)
    set_run_font(ref.add_run("Корпоративные компетенции: "), size=8, bold=True, color=NEAR_BLACK)
    set_run_font(ref.add_run(" · ".join(CORP_COMPETENCIES)), size=7.5, color=GRAY)

    h1 = doc.add_paragraph()
    clear_paragraph(h1)
    h1.paragraph_format.space_before = Pt(2)
    h1.paragraph_format.space_after = Pt(5)
    set_run_font(h1.add_run("1. Действия и корпоративные компетенции"), size=12, bold=True)
    add_horizontal_line(h1, color="26A6E0", size="16")

    action_table = doc.add_table(rows=1 + len(ACTIONS), cols=3)
    set_table_full_width(action_table)
    write_cell(
        action_table.rows[0].cells[0],
        "Действие / поведенческий индикатор\n(отдел продуктового запуска)",
        size=8.5,
        bold=True,
        fill=HEADER_BG,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    write_cell(
        action_table.rows[0].cells[1],
        "Корпоративные компетенции",
        size=8.5,
        bold=True,
        fill=HEADER_BG,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    write_cell(
        action_table.rows[0].cells[2],
        "Как у сотрудника в реальности",
        size=8.5,
        bold=True,
        fill=HEADER_BG,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    for i, (action, comps) in enumerate(ACTIONS):
        row = action_table.rows[i + 1]
        bg = WHITE if i % 2 == 0 else ROW_ALT
        write_cell(row.cells[0], f"{i + 1}. {action}", size=8, fill=bg)
        write_cell(row.cells[1], format_comps(comps), size=7.5, fill=COMP_BG)
        write_cell(row.cells[2], "", size=8, fill=REALITY_BG)
        row.height = Cm(1.15)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_col_widths(action_table, [11.5, 8.5, 7.0])

    h2 = doc.add_paragraph()
    clear_paragraph(h2)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(5)
    set_run_font(h2.add_run("2. Ожидания по работе"), size=12, bold=True)
    add_horizontal_line(h2, color="26A6E0", size="16")

    exp_table = doc.add_table(rows=1 + len(EXPECTATIONS), cols=3)
    set_table_full_width(exp_table)
    write_cell(
        exp_table.rows[0].cells[0],
        "Ожидание по работе",
        size=8.5,
        bold=True,
        fill=HEADER_BG,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    write_cell(
        exp_table.rows[0].cells[1],
        "Корпоративные компетенции",
        size=8.5,
        bold=True,
        fill=HEADER_BG,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    write_cell(
        exp_table.rows[0].cells[2],
        "Реальность",
        size=8.5,
        bold=True,
        fill=HEADER_BG,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    for i, (expectation, comps) in enumerate(EXPECTATIONS):
        row = exp_table.rows[i + 1]
        bg = WHITE if i % 2 == 0 else ROW_ALT
        write_cell(row.cells[0], f"{i + 1}. {expectation}", size=8, fill=bg)
        write_cell(row.cells[1], format_comps(comps), size=7.5, fill=COMP_BG)
        write_cell(row.cells[2], "", size=8, fill=REALITY_BG)
        row.height = Cm(1.05)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_col_widths(exp_table, [11.5, 8.5, 7.0])

    h3 = doc.add_paragraph()
    clear_paragraph(h3)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(5)
    set_run_font(h3.add_run("3. Итог и договорённости"), size=12, bold=True)
    add_horizontal_line(h3, color="26A6E0", size="16")

    summary = doc.add_table(rows=4, cols=2)
    set_table_full_width(summary)
    summary_items = [
        ("Сильные стороны (какие компетенции проявлены сильнее)", ""),
        ("Зоны роста (приоритетные компетенции и действия)", ""),
        ("Конкретные договорённости / план развития", ""),
        (
            "Общая оценка соответствия роли",
            "☐ не соответствует   ☐ частично   ☐ соответствует   ☐ превосходит",
        ),
    ]
    for i, (label, value) in enumerate(summary_items):
        write_cell(summary.rows[i].cells[0], label, size=8.5, bold=True, fill=LABEL_BG)
        write_cell(summary.rows[i].cells[1], value, size=8.5, fill=WHITE)
        summary.rows[i].height = Cm(1.4)
        summary.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    set_col_widths(summary, [8.0, 19.0])

    sig_title = doc.add_paragraph()
    clear_paragraph(sig_title)
    sig_title.paragraph_format.space_before = Pt(12)
    sig_title.paragraph_format.space_after = Pt(6)
    set_run_font(sig_title.add_run("Подписи"), size=11, bold=True)

    sig = doc.add_table(rows=2, cols=2)
    set_table_full_width(sig)
    write_cell(
        sig.rows[0].cells[0],
        "Оценщик / руководитель:\n\n________________ / ________________",
        size=9,
        fill=WHITE,
    )
    write_cell(
        sig.rows[0].cells[1],
        "Сотрудник:\n\n________________ / ________________",
        size=9,
        fill=WHITE,
    )
    write_cell(
        sig.rows[1].cells[0],
        "Дата: «____» ______________ 20____ г.",
        size=8,
        color=GRAY,
        fill=WHITE,
    )
    write_cell(
        sig.rows[1].cells[1],
        "Дата: «____» ______________ 20____ г.",
        size=8,
        color=GRAY,
        fill=WHITE,
    )
    set_col_widths(sig, [13.5, 13.5])
    for row in sig.rows:
        row.height = Cm(1.6)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    foot = doc.add_paragraph()
    clear_paragraph(foot)
    foot.paragraph_format.space_before = Pt(10)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        foot.add_run(
            "© ГК Форус  ·  Внутренний документ отдела продуктового запуска  ·  "
            "Шаблон для индивидуального анализа сотрудника"
        ),
        size=7,
        color=GRAY,
    )

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    ART_RU.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    doc.save(OUT_ROOT)
    doc.save(ART_RU)
    doc.save(ART_EN)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Saved: {path}")
