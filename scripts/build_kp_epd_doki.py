#!/usr/bin/env python3
"""Коммерческое предложение: 1С-ЭПД и Доки.Логистика (стиль ГК Форус).

Акцент на Доки.Логистика, понятный язык без сокращений для бухгалтеров и клиентов.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "КП_1С-ЭПД_и_Доки_Логистика.docx"
ASSETS = ROOT / "assets_forus"

YELLOW = RGBColor(0xE8, 0xB8, 0x4A)
YELLOW_HEX = "E8B84A"
DARK = RGBColor(0x1E, 0x1E, 0x1E)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = "F5F5F5"
WHITE = "FFFFFF"
SOFT_YELLOW = "FFF8E7"
SOFT_DOKI = "F3F0FA"
HEADER_BG = "1E1E1E"
ACCENT_DOKI = "2D2A4A"


def set_run_font(run, name="Arial", size=10, bold=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_table_borders(table, color="CCCCCC", sz="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def yellow_rule(doc, space_before=6, space_after=8, sz="24"):
    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(space_before)
    line.paragraph_format.space_after = Pt(space_after)
    pPr = line._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), YELLOW_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def clear_cell(cell):
    for p in cell.paragraphs:
        p.clear()


def section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, color=DARK)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), YELLOW_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def fill_header_cell(cell, text, center=True, bg=HEADER_BG):
    clear_cell(cell)
    set_cell_shading(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))


def fill_cell(cell, text, *, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, shade=None, color=DARK):
    clear_cell(cell)
    if shade:
        set_cell_shading(cell, shade)
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_body(doc, text, *, size=10, bold=False, color=DARK, space_after=6, space_before=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_rich(doc, parts, *, size=10, space_after=6, space_before=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    """parts: list of (text, bold, color)."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.2
    for text, bold, color in parts:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_check(cell_or_doc, text, *, size=9, first=False):
    if hasattr(cell_or_doc, "rows"):
        raise TypeError("pass cell or doc")
    if hasattr(cell_or_doc, "paragraphs") and not hasattr(cell_or_doc, "add_heading"):
        # cell
        p = cell_or_doc.paragraphs[0] if first else cell_or_doc.add_paragraph()
        if first:
            clear_cell(cell_or_doc)
            p = cell_or_doc.paragraphs[0]
    else:
        p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    m = p.add_run("●  ")
    set_run_font(m, size=size, color=YELLOW, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=size, color=DARK)
    return p


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)

    # --- Шапка ---
    header_table = doc.add_table(rows=1, cols=2)
    left, right = header_table.rows[0].cells
    clear_cell(left)
    clear_cell(right)

    logo_path = ASSETS / "brand" / "forus_logo_word.png"
    if logo_path.exists():
        run = left.paragraphs[0].add_run()
        run.add_picture(str(logo_path), width=Cm(4.0))
    else:
        run = left.paragraphs[0].add_run("Форус")
        set_run_font(run, size=22, bold=True)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, (line, bold) in enumerate(
        [
            ("Группа компаний «Форус»", True),
            ("Крупнейшая IT-компания Иркутской области", False),
            ("г. Иркутск, ул. Ямская, 1/1", False),
            ("+7 (3952) 78-00-00  ·  www.forus.ru", False),
        ]
    ):
        if i:
            p = right.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, size=9, bold=bold, color=DARK if bold else GRAY)

    yellow_rule(doc, space_before=4, space_after=8)

    # --- Заголовок ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    set_run_font(run, size=17, bold=True, color=DARK)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    run = sub.add_run("Переход на электронные перевозочные документы")
    set_run_font(run, size=11, color=GRAY)

    products = doc.add_paragraph()
    products.alignment = WD_ALIGN_PARAGRAPH.CENTER
    products.paragraph_format.space_after = Pt(8)
    run = products.add_run("Доки.Логистика")
    set_run_font(run, size=13, bold=True, color=DARK)
    run = products.add_run("  и  ")
    set_run_font(run, size=12, color=GRAY)
    run = products.add_run("1С-ЭПД")
    set_run_font(run, size=12, bold=True, color=GRAY)

    # --- Обращение ---
    add_body(
        doc,
        "Уважаемые коллеги!",
        size=11,
        bold=True,
        space_after=4,
    )
    add_body(
        doc,
        "Группа компаний «Форус» благодарит вас за интерес к нашим решениям и готовность "
        "выстроить удобный электронный документооборот по грузоперевозкам.",
        size=10,
        space_after=6,
    )
    add_body(
        doc,
        "С 1 сентября 2026 года вступает в силу Федеральный закон от 07.06.2025 № 140-ФЗ. "
        "Организации, которые участвуют в автомобильных грузоперевозках, переходят на "
        "электронные перевозочные документы: электронную транспортную накладную, "
        "электронный заказ-заявку, экспедиторские документы и электронный путевой лист "
        "(в установленных законом случаях). Бумажный обмен по этим документам больше "
        "не будет соответствовать требованиям законодательства.",
        size=10,
        space_after=6,
    )
    add_rich(
        doc,
        [
            (
                "Мы поможем вашей компании подключиться вовремя, без лишней нагрузки на бухгалтерию "
                "и логистику. Ниже — два решения, которые мы предлагаем. Наше приоритетное "
                "рекомендование для большинства клиентов — сервис ",
                False,
                DARK,
            ),
            ("Доки.Логистика", True, DARK),
            (
                ": он удобнее в ежедневной работе, гибче по сценариям и позволяет вести "
                "весь электронный документооборот в одном окне.",
                False,
                DARK,
            ),
        ],
        size=10,
        space_after=8,
    )

    # --- Рекомендуем Доки ---
    section_title(doc, "Рекомендуем: Доки.Логистика")

    recommend = doc.add_table(rows=1, cols=1)
    set_table_borders(recommend, color=YELLOW_HEX, sz="12")
    box = recommend.rows[0].cells[0]
    set_cell_shading(box, SOFT_YELLOW)
    clear_cell(box)
    p = box.paragraphs[0]
    run = p.add_run("Почему мы советуем начать именно с Доки.Логистика")
    set_run_font(run, size=11, bold=True, color=DARK)

    recommend_points = [
        "Сервис работает сразу в трёх средах: в программе 1С, в веб-кабинете через браузер и в мобильном приложении. Можно пользоваться всеми вариантами сразу или выбрать один — например, только телефон или только браузер.",
        "Подходит компаниям без 1С, с редко обновляемой 1С, а также тем, у кого логист, водитель или руководитель склада не работают в учётной базе бухгалтера.",
        "В одном сервисе — и перевозочные документы, и обычный электронный документооборот с контрагентами: счета, акты, универсальные передаточные документы, договоры.",
        "Документы хранятся в защищённом облачном архиве. Даже если с компьютером или базой 1С что-то случится, история обмена не потеряется.",
        "Для новых клиентов действует промотариф: 3 месяца работы без ограничений по количеству отправок (при подключении пакета документов).",
        "Водитель может подписать документы с телефона: простой электронной подписью, усиленной квалифицированной электронной подписью или через Госключ.",
    ]
    for t in recommend_points:
        p = box.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(2)
        m = p.add_run("●  ")
        set_run_font(m, size=9, color=YELLOW, bold=True)
        r = p.add_run(t)
        set_run_font(r, size=9, color=DARK)

    add_body(
        doc,
        "Доки.Логистика — это сервис группы компаний «Астрал», надёжного оператора "
        "электронного документооборота. Подключение и сопровождение для вас выполняет "
        "ГК «Форус»: настройка, обучение, поддержка и подбор тарифа под ваш объём документов.",
        size=9,
        color=GRAY,
        space_before=6,
        space_after=4,
    )

    # --- Подробные плюсы Доки ---
    section_title(doc, "Доки.Логистика — возможности для вашей компании")

    doki_blocks = [
        (
            "Удобно бухгалтеру и логисту",
            [
                "Из 1С документы уходят в электронный обмен в один клик — без выгрузки файлов и ручной пересылки.",
                "Входящие документы можно сразу превратить в учётные документы 1С: меньше расхождений в реквизитах и ручного ввода.",
                "Остаток пакета и статус тарифа видны прямо в 1С и в веб-кабинете.",
                "Руководитель может подписать документ из дома, командировки или с телефона — бухгалтеру не нужно ждать его у рабочего компьютера.",
            ],
        ),
        (
            "Гибкая работа без привязки к одному компьютеру",
            [
                "Веб-кабинет синхронизирован с 1С: статусы документов обновляются сразу в обоих интерфейсах.",
                "Мобильное приложение позволяет водителю подтвердить погрузку и выгрузку на месте.",
                "Можно разделить доступ по сотрудникам и подразделениям: бухгалтер, логист и склад работают в своих зонах ответственности.",
                "Подходит, если бухгалтер не готов пускать логиста в основную базу 1С — обмен ведётся в отдельном удобном контуре.",
            ],
        ),
        (
            "Экономия и прозрачность",
            [
                "Отправка электронного документа в несколько раз дешевле печати, конверта и курьерской доставки (для сравнения: бумажный документ часто обходится около 50 рублей).",
                "Статус документа всегда на виду: отправлен, получен, подписан, требуется действие.",
                "Быстрее закрываете период и ускоряете получение оплаты за счёт мгновенной доставки закрывающих документов.",
                "Чем больше пакет документов — тем ниже цена одной отправки. Неиспользованный остаток переносится при своевременном продлении.",
            ],
        ),
    ]

    for title_text, items in doki_blocks:
        add_body(doc, title_text, size=10, bold=True, space_before=4, space_after=3)
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.2)
            m = p.add_run("●  ")
            set_run_font(m, size=9, color=YELLOW, bold=True)
            r = p.add_run(item)
            set_run_font(r, size=9, color=DARK)

    # --- 1С-ЭПД кратко ---
    section_title(doc, "Альтернатива: сервис 1С-ЭПД")

    add_body(
        doc,
        "1С-ЭПД — типовое решение фирмы «1С» для обмена электронными перевозочными "
        "документами внутри программ 1С. Оно хорошо подходит компаниям, которые уже "
        "активно работают в актуальной типовой 1С, регулярно обновляются и хотят вести "
        "перевозочный обмен прямо в учётной системе без отдельного веб-кабинета.",
        size=10,
        space_after=4,
    )

    epd_points = [
        "Встроено в типовые конфигурации 1С — отдельное расширение обычно не требуется.",
        "Поддерживает форматы электронных перевозочных документов, которые требует налоговая служба.",
        "Есть обмен с другими операторами электронных перевозочных документов.",
        "Водитель может подписывать документы в мобильном приложении усиленной квалифицированной электронной подписью.",
        "При покупке пакета от 1 000 документов действует акция: бесплатная настройка одного рабочего места.",
    ]
    for t in epd_points:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        m = p.add_run("●  ")
        set_run_font(m, size=9, color=GRAY, bold=True)
        r = p.add_run(t)
        set_run_font(r, size=9, color=DARK)

    add_body(
        doc,
        "Ограничение: мобильное приложение 1С-ЭПД сейчас работает в связке с 1С. "
        "Если у сотрудников нет постоянного доступа к базе 1С, им нужен только телефон "
        "или браузер, либо вы хотите совместить перевозочные и обычные документы "
        "в одном сервисе — удобнее выбрать Доки.Логистика.",
        size=9,
        color=GRAY,
        space_before=4,
        space_after=4,
    )

    # --- Сравнение (в пользу Доки) ---
    section_title(doc, "Сравнение по возможностям")

    add_body(
        doc,
        "Ниже — наглядное сравнение. Мы специально выделили параметры, которые важны "
        "бухгалтерии, логистике и руководству в повседневной работе.",
        size=9,
        color=GRAY,
        space_after=4,
    )

    diff = doc.add_table(rows=1, cols=3)
    set_table_borders(diff, color="DDDDDD", sz="4")
    fill_header_cell(diff.rows[0].cells[0], "Что важно клиенту", center=False)
    fill_header_cell(diff.rows[0].cells[1], "Доки.Логистика")
    fill_header_cell(diff.rows[0].cells[2], "1С-ЭПД")

    rows_data = [
        (
            "Где можно работать",
            "В 1С, в браузере и в мобильном приложении — независимо друг от друга",
            "В основном внутри программы 1С (+ мобильное приложение в связке с 1С)",
        ),
        (
            "Нужна ли 1С для старта",
            "Не обязательна: можно начать с веб-кабинета или телефона",
            "Нужна база 1С (или бесплатная программа «1С:Клиент ЭДО» для обмена документами)",
        ),
        (
            "Кто может работать в сервисе",
            "Бухгалтер, логист, склад, руководитель, водитель — с разграничением прав",
            "Пользователи, у которых есть доступ к базе 1С",
        ),
        (
            "Какие документы доступны",
            "Перевозочные документы + полный электронный документооборот с контрагентами",
            "Фокус на электронных перевозочных документах",
        ),
        (
            "Где хранятся документы",
            "Надёжный облачный архив + синхронизация с 1С",
            "В информационной базе 1С",
        ),
        (
            "Подписание вне офиса",
            "Удобно из браузера и с телефона без доступа к рабочему компьютеру",
            "Возможно через мобильное приложение при настроенной связке с 1С",
        ),
        (
            "Сценарий «логист отдельно от бухгалтера»",
            "Да: логист работает в сервисе, не заходя в основную базу 1С",
            "Как правило, работа идёт через учётную систему 1С",
        ),
        (
            "Старт для новых клиентов",
            "Промотариф: 3 месяца безлимитных отправок при покупке пакета",
            "Акция: бесплатная настройка рабочего места при пакете от 1 000 документов",
        ),
        (
            "Ежедневное удобство",
            "Одно окно для перевозок и обычного обмена документами",
            "Отдельный контур для перевозочных документов в 1С",
        ),
    ]
    for i, (a, b, c) in enumerate(rows_data):
        row = diff.add_row().cells
        shade = SOFT_DOKI if i % 2 == 0 else WHITE
        fill_cell(row[0], a, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, shade=shade)
        fill_cell(row[1], b, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, shade=shade)
        fill_cell(row[2], c, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, shade=LIGHT_GRAY if i % 2 == 0 else WHITE)

    add_body(
        doc,
        "Итог для большинства компаний: Доки.Логистика даёт больше свободы в организации "
        "работы и закрывает и перевозки, и обычный электронный документооборот. "
        "1С-ЭПД оставляем как вариант для тех, кто хочет остаться строго внутри типовой 1С.",
        size=9,
        bold=True,
        space_before=6,
        space_after=6,
    )

    # --- Тарифы Доки ---
    section_title(doc, "Стоимость пакетов документов Доки.Логистика")

    add_body(
        doc,
        "Вы оплачиваете годовой пакет исходящих документов. Входящие документы от "
        "контрагентов можно получать и подписывать без покупки тарифа. Чем больше пакет — "
        "тем ниже цена одной отправки. Срок действия тарифа — 12 месяцев. "
        "Неизрасходованный остаток переносится на следующий период, если новый пакет "
        "куплен вовремя и стоит в очереди на активацию.",
        size=9,
        space_after=4,
    )

    t_doki = doc.add_table(rows=1, cols=3)
    set_table_borders(t_doki, color="DDDDDD", sz="4")
    fill_header_cell(t_doki.rows[0].cells[0], "Количество документов в год", center=False)
    fill_header_cell(t_doki.rows[0].cells[1], "Цена за 1 документ, руб.")
    fill_header_cell(t_doki.rows[0].cells[2], "Стоимость тарифа, руб. / год")

    doki_packages = [
        ("200", "9,00", "1 800"),
        ("500", "8,40", "4 200"),
        ("1 200", "7,67", "9 200"),
        ("1 500", "7,60", "11 400"),
        ("3 000", "5,70", "17 100"),
        ("5 000", "4,90", "24 500"),
        ("10 000", "3,90", "39 000"),
        ("20 000", "3,65", "73 000"),
        ("30 000", "3,50", "105 000"),
        ("50 000", "3,40", "170 000"),
        ("100 000", "3,20", "320 000"),
    ]
    for i, (qty, per, total) in enumerate(doki_packages):
        row = t_doki.add_row().cells
        shade = SOFT_DOKI if i % 2 == 0 else WHITE
        fill_cell(row[0], qty, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, shade=shade)
        fill_cell(row[1], per, size=9, shade=shade)
        fill_cell(row[2], total, bold=True, size=9, shade=shade)

    promo = doc.add_table(rows=1, cols=1)
    set_table_borders(promo, color=YELLOW_HEX, sz="10")
    pc = promo.rows[0].cells[0]
    set_cell_shading(pc, SOFT_YELLOW)
    clear_cell(pc)
    p = pc.paragraphs[0]
    run = p.add_run("Акция «15 месяцев по цене 12»")
    set_run_font(run, size=10, bold=True, color=DARK)
    p = pc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(
        "При покупке любого пакета Доки и 1 часа линии консультаций за 3 660 рублей "
        "мы даём дополнительно 3 месяца демо-доступа. Сначала вы пользуетесь демо-периодом, "
        "затем автоматически активируется оплаченный пакет. Итого — 15 месяцев работы "
        "вместо 12. Промотариф на 3 месяца безлимита доступен новым клиентам "
        "(не применяется, если у вас уже есть учётная запись в связанном сервисе Астрал.ЭДО)."
    )
    set_run_font(run, size=9, color=DARK)

    # --- Тарифы 1С-ЭПД ---
    section_title(doc, "Стоимость пакетов документов 1С-ЭПД")

    add_body(
        doc,
        "Для 1С-ЭПД доступны предоплатные пакеты на 12 месяцев и постоплатная модель "
        "(7 рублей за один отправленный титул документа). "
        "Важные правила оплаты титулов: в электронном заказе-заявке титулы бесплатны "
        "для грузоотправителя и перевозчика; в электронной транспортной накладной "
        "оплачиваются титул грузоотправителя и титул перевозчика; в электронном "
        "путевом листе оплачивается один титул. Правила актуальны до 31.12.2026.",
        size=9,
        space_after=4,
    )

    t_epd = doc.add_table(rows=1, cols=3)
    set_table_borders(t_epd, color="DDDDDD", sz="4")
    fill_header_cell(t_epd.rows[0].cells[0], "Пакет на 12 месяцев", center=False)
    fill_header_cell(t_epd.rows[0].cells[1], "Цена за 1 документ, руб.")
    fill_header_cell(t_epd.rows[0].cells[2], "Стоимость пакета, руб.")

    epd_packages = [
        ("Пакет на 600 документов", "6,00", "3 600"),
        ("Пакет на 1 000 документов", "5,00", "5 000"),
        ("Пакет на 5 000 документов", "4,50", "22 500"),
        ("Пакет на 10 000 документов", "4,00", "40 000"),
        ("Пакет на 50 000 документов", "3,00", "150 000"),
        ("Пакет на 100 000 документов", "2,50", "250 000"),
    ]
    for i, (name, per, total) in enumerate(epd_packages):
        row = t_epd.add_row().cells
        shade = SOFT_YELLOW if i % 2 == 0 else WHITE
        fill_cell(row[0], name, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, shade=shade)
        fill_cell(row[1], per, size=9, shade=shade)
        fill_cell(row[2], total, bold=True, size=9, shade=shade)

    add_body(
        doc,
        "Акция по 1С-ЭПД: при покупке пакета от 1 000 документов — бесплатная настройка "
        "одного рабочего места (активация сервиса, КриптоПро, электронная подпись, "
        "машиночитаемая доверенность, приглашение до трёх контрагентов). "
        "Лицензия средства криптозащиты и сертификаты электронной подписи оплачиваются отдельно.",
        size=9,
        space_before=4,
        space_after=6,
    )

    # --- Дополнительно ---
    section_title(doc, "Что может понадобиться дополнительно")

    extra = doc.add_table(rows=1, cols=3)
    set_table_borders(extra, color="DDDDDD", sz="4")
    fill_header_cell(extra.rows[0].cells[0], "Услуга или продукт", center=False)
    fill_header_cell(extra.rows[0].cells[1], "Стоимость")
    fill_header_cell(extra.rows[0].cells[2], "Для чего нужно", center=False)

    extras = [
        (
            "Линия консультаций 1С (от 1 до 3 часов)",
            "3 660 руб./час",
            "Удалённая помощь по настройке и вопросам работы",
        ),
        (
            "Подготовка одного рабочего места под электронные перевозочные документы",
            "7 320 руб.",
            "Обучение, экспресс-анализ сценариев, базовая настройка",
        ),
        (
            "Лицензия КриптоПро CSP (бессрочно, 1 рабочее место)",
            "3 700 руб.",
            "Программа для работы с электронной подписью",
        ),
        (
            "Усиленная квалифицированная электронная подпись",
            "1 050 руб.",
            "Выпускается на каждого сотрудника, который подписывает документы",
        ),
        (
            "Рутокен 3.0 или NFC-токен для подписи",
            "от 2 700 руб.",
            "Носитель подписи; NFC нужен для подписания с телефона",
        ),
        (
            "Договор 1С:ИТС",
            "от 3 273 руб./мес.",
            "Сопровождение и обновление программ 1С",
        ),
    ]
    for i, (n, price, c) in enumerate(extras):
        row = extra.add_row().cells
        shade = LIGHT_GRAY if i % 2 == 0 else WHITE
        fill_cell(row[0], n, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, shade=shade)
        fill_cell(row[1], price, size=8, shade=shade)
        fill_cell(row[2], c, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, shade=shade)

    # --- О компании Форус ---
    section_title(doc, "Почему Форус")

    about = doc.add_table(rows=1, cols=1)
    set_table_borders(about, color=YELLOW_HEX, sz="12")
    ac = about.rows[0].cells[0]
    set_cell_shading(ac, SOFT_YELLOW)
    clear_cell(ac)
    p = ac.paragraphs[0]
    run = p.add_run("Группа компаний «Форус»")
    set_run_font(run, size=11, bold=True, color=DARK)
    p = ac.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    run = p.add_run(
        "Крупнейшая IT-компания Иркутской области, на рынке с 1992 года. "
        "Входим в ТОП-50 крупнейших IT-компаний России. В штате более 300 "
        "сертифицированных специалистов. Первые в рейтинге фирмы «1С» по числу "
        "аттестованных специалистов по «1С:Предприятие 8» в Иркутской области. "
        "Система менеджмента качества соответствует международному стандарту ISO 9001."
    )
    set_run_font(run, size=9, color=DARK)

    statuses = doc.add_table(rows=2, cols=3)
    set_table_borders(statuses, color="DDDDDD", sz="4")
    status_items = [
        "1С: Центр ERP",
        "Центр сопровождения 1С",
        "Центр компетенции 1С: КОРП",
        "Центр реальной автоматизации",
        "Центр компетенции по документообороту",
        "Центр компетенции по кадровому электронному документообороту",
    ]
    # fill 2x3
    for idx, text in enumerate(status_items):
        r, c = divmod(idx, 3)
        cell = statuses.rows[r].cells[c]
        set_cell_shading(cell, WHITE)
        clear_cell(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        m = p.add_run("● ")
        set_run_font(m, size=8, color=YELLOW, bold=True)
        run = p.add_run(text)
        set_run_font(run, size=8, bold=True, color=DARK)

    add_body(
        doc,
        "Мы не просто продаём доступ к сервису: помогаем выбрать решение под ваши процессы, "
        "настраиваем рабочее место, выпускаем электронные подписи, обучаем сотрудников "
        "и сопровождаем после запуска. Вы всегда можете обратиться к персональному менеджеру "
        "и на линию консультаций Форус.",
        size=9,
        space_before=6,
        space_after=6,
    )

    # --- Следующие шаги ---
    section_title(doc, "Как начать работу")

    steps = [
        "Короткий созвон или встреча: уточняем ваши сценарии перевозок и объём документов.",
        "Подбираем пакет Доки.Логистика (или 1С-ЭПД, если это ваш осознанный выбор) и считаем итоговую стоимость.",
        "Подключаем сервис, настраиваем подписи и права доступа, приглашаем ключевых контрагентов.",
        "Обучаем сотрудников и сопровождаем на старте, чтобы обмен пошёл без сбоев.",
    ]
    for i, t in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        n = p.add_run(f"{i}. ")
        set_run_font(n, size=10, bold=True, color=YELLOW)
        r = p.add_run(t)
        set_run_font(r, size=10, color=DARK)

    # --- Контакты ---
    yellow_rule(doc, space_before=10, space_after=8)

    contact_box = doc.add_table(rows=1, cols=2)
    set_table_borders(contact_box, color=YELLOW_HEX, sz="12")
    left_c, right_c = contact_box.rows[0].cells
    set_cell_shading(left_c, SOFT_YELLOW)
    set_cell_shading(right_c, SOFT_YELLOW)

    clear_cell(left_c)
    p = left_c.paragraphs[0]
    run = p.add_run("Ваш менеджер")
    set_run_font(run, size=9, color=GRAY)
    p = left_c.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Оглоблина Софья")
    set_run_font(run, size=13, bold=True, color=DARK)
    p = left_c.add_paragraph()
    run = p.add_run("Менеджер по продаже решений для электронных перевозочных документов")
    set_run_font(run, size=8, color=DARK)
    for line in [
        "Электронная почта: sogloblina@forus.ru",
        "Телефон: +7 (3952) 78-00-00, доб. 1861",
        "Разница со временем Москвы: +5 часов",
    ]:
        p = left_c.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, size=9, color=DARK)

    clear_cell(right_c)
    p = right_c.paragraphs[0]
    run = p.add_run("Контакты группы компаний «Форус»")
    set_run_font(run, size=9, color=GRAY)
    p = right_c.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("ООО НПФ «Форус»")
    set_run_font(run, size=12, bold=True, color=DARK)
    for line in [
        "664047, г. Иркутск, ул. Ямская, 1/1, офис 1",
        "Телефон: +7 (3952) 78-00-00, 72-87-02",
        "Электронная почта: info@forus.ru",
        "Сайт: www.forus.ru",
        "ИНН 3812023430  ·  ОГРН 1023801752633",
    ]:
        p = right_c.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, size=9, color=DARK)

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(12)
    run = closing.add_run(
        "Готовы подобрать пакет под ваш объём документов и провести демонстрацию Доки.Логистика.\n"
        "Будем рады ответить на вопросы и помочь с быстрым запуском."
    )
    set_run_font(run, size=9, color=GRAY)

    doc.save(OUT)
    root_copy = ROOT / "КП_1С-ЭПД_и_Доки_Логистика.docx"
    doc.save(root_copy)
    print(f"Saved: {OUT}")
    print(f"Saved: {root_copy}")
    return OUT


if __name__ == "__main__":
    from PIL import Image, ImageDraw, ImageFont

    brand = ASSETS / "brand"
    brand.mkdir(parents=True, exist_ok=True)
    logo_path = brand / "forus_logo_word.png"
    if not logo_path.exists():
        font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
        img = Image.new("RGBA", (420, 120), (255, 255, 255, 255))
        d = ImageDraw.Draw(img)
        font = ImageFont.truetype(font_path, 64)
        text = "Форус"
        x, y = 10, 30
        d.text((x, y), text, fill=(30, 30, 30, 255), font=font)
        bbox_f = d.textbbox((0, 0), "Ф", font=font)
        fw = bbox_f[2] - bbox_f[0]
        bbox_o = d.textbbox((0, 0), "о", font=font)
        ow = bbox_o[2] - bbox_o[0]
        ox = x + fw + 2
        oy = y - 6
        d.rounded_rectangle([ox, oy, ox + ow - 2, oy + 9], radius=2, fill=(232, 184, 74, 255))
        img.convert("RGB").save(logo_path)

    build()
