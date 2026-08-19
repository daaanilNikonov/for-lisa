#!/usr/bin/env python3
"""Generate Forus-branded job description for Product Launch Team Lead."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Должностная_инструкция_тимлид_продуктовый_запуск.docx"
OUT_ROOT = ROOT / "Должностная_инструкция_тимлид_продуктовый_запуск.docx"
ART_RU = Path("/opt/cursor/artifacts/Должностная_инструкция_тимлид_продуктовый_запуск.docx")
ART_EN = Path("/opt/cursor/artifacts/Dolzhnostnaya_instrukciya_timlid_produktovyy_zapusk.docx")

YELLOW = RGBColor(0xFC, 0xCD, 0x68)
NEAR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x76, 0x76, 0x77)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LABEL_BG = RGBColor(0xFF, 0xF3, 0xD6)
ROW_ALT = RGBColor(0xFF, 0xF8, 0xE8)
HEADER_BG = YELLOW
PRIORITY_BG = RGBColor(0xF0, 0xF7, 0xFB)
FONT = "Verdana"

# Ранг 1 = сильнее влияет на фактическую нагрузку роли.
# Норма месяца ≈ 168–176 ч. Сумма ориентиров ≈ 171 ч.
# Ядро: стратегии / аналитика / финансы / продвижение / встречи (~75%+); МПП — обязательный контур.
FUNCTIONS_TABLE: list[dict] = [
    {
        "rank": 1,
        "function": "Проработка продуктовых стратегий по портфелю запусков (одновременно 3–5+ продуктов): гипотезы, позиционирование, приоритеты вывода на рынок",
        "hours": "36 ч / мес (~21%)",
        "result": "Стратегия / one-pager по каждому продукту; приоритеты квартала; go / iterate / stop; обновлённый оффер",
        "needs": "Карточки продуктов; ограничения ЦРП; конкурентный контекст; факты воронки; решения руководства по приоритетам портфеля",
    },
    {
        "rank": 2,
        "function": "Аналитика базы: сегменты, качество контактов, конверсии этапов, «узкие места» воронки по продуктам и МПП",
        "hours": "28 ч / мес (~16%)",
        "result": "Аналитические срезы базы; выводы по сегментам; список действий для МПП / маркетинга; корректировка скрипта и приоритетов обзвона",
        "needs": "Выгрузки / дашборды CRM; правила сегментации; история касаний; статусы сделок; при необходимости сверка с УТ",
    },
    {
        "rank": 3,
        "function": "Финансовая аналитика запусков: план/факт, экономика продуктов, мотивация, бюджеты и отклонения",
        "hours": "24 ч / мес (~14%)",
        "result": "Финмодель / план-факт; выводы по марже и окупаемости; корректировка KPI и мотивации; аргументы для решений по портфелю",
        "needs": "Выгрузки продаж из УТ; оплаты / отгрузки; правила мотивации; бюджеты рекламы; план продаж по продуктам",
    },
    {
        "rank": 4,
        "function": "Маркетинговые стратегии и организация мероприятий по продвижению сервисов (вебинары, акции, спецпредложения, партнёрские активности)",
        "hours": "22 ч / мес (~13%)",
        "result": "План продвижения; ТЗ / сценарий мероприятия; оффер акции; пост-аналитика эффекта на базу и сделки",
        "needs": "Слоты и ресурсы маркетинга; сегмент базы; креативы / лендинги; согласование формулировок; каналы регистрации и CRM-метки",
    },
    {
        "rank": 5,
        "function": "Встречи и синхронизация с другими подразделениями по продуктам (продукт / ЦРП, маркетинг, финблок, поддержка, смежные продажи)",
        "hours": "22 ч / мес (~13%)",
        "result": "Протоколы договорённостей; закрытые эскалации; согласованные сроки / ограничения продукта; единый статус для команды",
        "needs": "Календарь стейкхолдеров; статусы доработок; вводные по лицензиям / ограничениям; материалы запусков",
    },
    {
        "rank": 6,
        "function": "Операционная работа с МПП: канбаны, ключевые сделки, next steps и контроль исполнения по продуктам портфеля",
        "hours": "14 ч / мес (~8%)",
        "result": "Актуальная картина канбанов; список критичных сделок; решения по следующим шагам; прогноз по плану",
        "needs": "CRM / канбан МПП; стадии воронки; комментарии по сделкам; понимание текущего приоритета продуктов",
    },
    {
        "rank": 7,
        "function": "Контроль качества продаж и короткое обучение МПП под текущие продукты / акции / изменения оффера",
        "hours": "8 ч / мес (~5%)",
        "result": "Обновлённые скрипты / шпаргалки; мини-разборы; быстрый допуск МПП к изменениям продукта",
        "needs": "Актуальный оффер; записи звонков; краткие product notes; критерии качества",
    },
    {
        "rank": 8,
        "function": "Подготовка КП, скриптов и продуктовых материалов под несколько параллельных запусков",
        "hours": "6 ч / мес (~4%)",
        "result": "КП / one-pager / скрипт / FAQ по продуктам портфеля",
        "needs": "Вводные ЦРП; согласованные формулировки; конкурентные отличия; ограничения внедрения",
    },
    {
        "rank": 9,
        "function": "1-1, мотивация и кадровые решения по МПП",
        "hours": "4 ч / мес (~2%)",
        "result": "Протоколы 1-1; решения по мотивации / рискам; точечные ИПР",
        "needs": "Факт KPI; данные по сделкам; правила мотивации",
    },
    {
        "rank": 10,
        "function": "Развитие партнёрского / агентского канала совместно с ЦРП",
        "hours": "3 ч / мес (~2%)",
        "result": "Статус партнёрского канала; план касаний; гипотезы через партнёров",
        "needs": "Условия партнёрок; контакты ЦРП; CRM партнёрских сделок",
    },
    {
        "rank": 11,
        "function": "Участие в подборе / онбординге МПП и поддержка процессов группы",
        "hours": "2 ч / мес (~1%)",
        "result": "Оценка кандидата / план адаптации; точечные правки регламентов",
        "needs": "Профиль роли; чек-лист онбординга; as-is процессов",
    },
    {
        "rank": 12,
        "function": "Собственное развитие и фиксация лучших практик запусков",
        "hours": "2 ч / мес (~1%)",
        "result": "Короткие практики / памятки; личные выводы по портфелю продуктов",
        "needs": "Доступ к обучению; база кейсов запусков",
    },
]

def set_run_font(run, size=10, bold=False, color=NEAR_BLACK, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), font)


def set_cell_shading(cell, color: RGBColor):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    tc_pr.append(shd)


def set_cell_borders(cell, color="D0D0D0", sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)
    tc_pr.append(borders)


def set_cell_margins(cell, top=40, bottom=40, left=50, right=50):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for name, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    for old in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(old)
    tc_pr.append(tc_mar)


def set_vertical_align(cell, val="center"):
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), val)
    for old in tc_pr.findall(qn("w:vAlign")):
        tc_pr.remove(old)
    tc_pr.append(v_align)


def clear_paragraph(paragraph):
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1


def write_cell(
    cell,
    text,
    size=8.5,
    bold=False,
    color=NEAR_BLACK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    fill=None,
):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if fill is not None:
        set_cell_shading(cell, fill)
    set_cell_borders(cell)
    set_cell_margins(cell)
    set_vertical_align(cell, "center")


def add_horizontal_line(paragraph, color="FCCD68", size="24"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_table_full_width(table):
    tbl_pr = table._tbl.tblPr
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")
    for old in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(old)
    tbl_pr.append(tbl_w)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_heading(doc, text):
    p = doc.add_paragraph()
    clear_paragraph(p)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run(text), size=12, bold=True)
    add_horizontal_line(p, color="26A6E0", size="16")


def add_body(doc, text, size=9.5, space_after=4, bold=False):
    p = doc.add_paragraph()
    clear_paragraph(p)
    p.paragraph_format.space_after = Pt(space_after)
    set_run_font(p.add_run(text), size=size, bold=bold)


def add_bullet(doc, text, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=size)
    else:
        set_run_font(p.add_run(text), size=size)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    # Landscape — таблица функций читается лучше
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    brand = doc.add_table(rows=1, cols=2)
    set_table_full_width(brand)
    write_cell(brand.rows[0].cells[0], "ФОРУС", size=16, bold=True, fill=YELLOW)
    write_cell(
        brand.rows[0].cells[1],
        "Группа компаний  ·  Отдел продуктового запуска",
        size=9,
        fill=YELLOW,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    set_col_widths(brand, [6.0, 21.0])

    title = doc.add_paragraph()
    clear_paragraph(title)
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ"), size=14, bold=True)
    add_horizontal_line(title, color="FCCD68", size="28")

    subtitle = doc.add_paragraph()
    clear_paragraph(subtitle)
    subtitle.paragraph_format.space_before = Pt(4)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(subtitle.add_run("Тимлид группы продуктового запуска"), size=12, bold=True)

    meta = doc.add_table(rows=4, cols=4)
    set_table_full_width(meta)
    meta_rows = [
        [("Должность", True), ("Тимлид группы продуктового запуска", False), ("Подразделение", True), ("Продуктовый запуск", False)],
        [("Категория", True), ("Руководитель", False), ("Подчинение", True), ("Руководитель направления", False)],
        [("В подчинении", True), ("Менеджеры по продажам (МПП)", False), ("Версия", True), ("1.2", False)],
        [("ФИО сотрудника", True), ("________________________________", False), ("Дата введения", True), ("«____» ________ 20____ г.", False)],
    ]
    for i, row_data in enumerate(meta_rows):
        for j, (text, is_label) in enumerate(row_data):
            cell = meta.rows[i].cells[j]
            if is_label:
                write_cell(cell, text, size=8, bold=True, fill=LABEL_BG)
            else:
                color = GRAY if "___" in text or text.startswith("«") else NEAR_BLACK
                write_cell(cell, text, size=8.5, color=color, fill=WHITE)
    set_col_widths(meta, [3.8, 9.7, 3.8, 9.7])

    note = doc.add_paragraph()
    clear_paragraph(note)
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(2)
    set_run_font(
        note.add_run(
            "Основание: карта функций / компетенций группы продуктового запуска (зона Тимлида). "
            "Приоритет роли — управление менеджерами и их сделками."
        ),
        size=8,
        color=GRAY,
    )

    add_heading(doc, "1. Общие положения")
    for t in [
        "1.1. Настоящая должностная инструкция определяет функциональные обязанности, права и ответственность Тимлида группы продуктового запуска ГК Форус.",
        "1.2. Тимлид относится к категории руководителей и назначается / освобождается от должности приказом руководителя в установленном порядке.",
        "1.3. Тимлид непосредственно подчиняется руководителю направления / куратору группы продуктового запуска.",
        "1.4. В непосредственном подчинении Тимлида находятся менеджеры по продажам (МПП) группы продуктового запуска.",
        "1.5. В своей работе Тимлид руководствуется законодательством РФ, локальными нормативными актами ГК Форус, регламентами взаимодействия подразделений, скриптами и стандартами продаж, а также настоящей инструкцией.",
    ]:
        add_body(doc, t)

    add_heading(doc, "2. Цель должности и принцип приоритетов")
    add_body(
        doc,
        "Фактический профиль роли Тимлида группы продуктового запуска — управление портфелем продуктов (одновременно минимум 3–5 запусков) "
        "через продуктовую стратегию, аналитику базы, финансовую аналитику, продвижение сервисов и плотную синхронизацию со смежными подразделениями. "
        "Работа с МПП (канбаны, сделки, обучение) остаётся обязательной зоной ответственности, но в текущей нагрузке занимает меньшую долю времени, "
        "чем продуктово-аналитический и кросс-функциональный контур.",
        bold=False,
    )
    add_body(doc, "Принцип распределения внимания (по фактической нагрузке):")
    for item in [
        "Основной объём времени (~75%+): продуктовые стратегии по 3–5+ продуктам, аналитика базы, финансовая аналитика, маркетинг / мероприятия продвижения, встречи со смежными подразделениями.",
        "Обязательный управленческий контур: канбаны и ключевые сделки МПП, короткое обучение под изменения продуктов, точечные 1-1.",
        "Сложность роли усиливается параллельным ведением нескольких продуктов: решения и встречи масштабируются не «на 1 продукт», а на весь портфель.",
        "Целевой баланс на перспективу: сохраняя сильный продуктово-аналитический контур, возвращать больше времени в работу с МПП и их сделками — это усиливает скорость проверки гипотез на рынке.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Матрица функций: приоритет, время, результат, входы")
    add_body(
        doc,
        "Функции проранжированы от наиболее влияющих на фактическую нагрузку роли к менее влияющим. "
        "Затраты времени — практический ориентир при норме ≈ 168–176 рабочих часов и параллельном ведении 3–5+ продуктов. "
        "Итого в таблице ≈ 171 ч; пиковые недели запусков / мероприятий могут превышать норму за счёт встреч и срочной аналитики.",
        size=8.5,
        space_after=6,
    )

    table = doc.add_table(rows=1 + len(FUNCTIONS_TABLE), cols=5)
    set_table_full_width(table)
    headers = [
        "Ранг\nважности",
        "Функция тимлида",
        "Затрата времени\n(ориентир / мес)",
        "Результат функции\n(документ / артефакт / показатель)",
        "Что требуется для выполнения\n(данные, доступы, входы)",
    ]
    for idx, header in enumerate(headers):
        write_cell(
            table.rows[0].cells[idx],
            header,
            size=8,
            bold=True,
            fill=HEADER_BG,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for i, item in enumerate(FUNCTIONS_TABLE):
        row = table.rows[i + 1]
        bg = WHITE if i % 2 == 0 else ROW_ALT
        # top ranks subtly highlighted
        rank_fill = PRIORITY_BG if item["rank"] <= 5 else bg
        write_cell(row.cells[0], str(item["rank"]), size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, fill=rank_fill)
        write_cell(row.cells[1], item["function"], size=7.5, fill=rank_fill if item["rank"] <= 5 else bg)
        write_cell(row.cells[2], item["hours"], size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER, fill=bg)
        write_cell(row.cells[3], item["result"], size=7.5, fill=bg)
        write_cell(row.cells[4], item["needs"], size=7.5, fill=bg)
        row.height = Cm(1.55)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    set_col_widths(table, [2.0, 7.2, 3.6, 7.0, 7.2])

    legend = doc.add_paragraph()
    clear_paragraph(legend)
    legend.paragraph_format.space_before = Pt(6)
    legend.paragraph_format.space_after = Pt(4)
    set_run_font(
        legend.add_run(
            "Подсветка рангов 1–5: фактическое ядро нагрузки — продуктовые стратегии, аналитика базы, финансы, "
            "продвижение / маркетинг и встречи со смежными подразделениями по портфелю продуктов."
        ),
        size=8,
        color=GRAY,
    )

    add_heading(doc, "4. Ключевые зоны ответственности (кратко)")
    add_body(doc, "4.1. Портфель продуктов и стратегия запуска (фактический приоритет №1)", bold=True)
    for item in [
        "Параллельно вести 3–5+ продуктов: для каждого — гипотеза, позиционирование, приоритеты и критерии успеха.",
        "Собирать продуктовую стратегию в решения, которые можно передать МПП и смежным отделам.",
        "Быстро пересобирать оффер / скрипт / приоритеты при изменении рынка или ограничений продукта.",
    ]:
        add_bullet(doc, item)

    add_body(doc, "4.2. Аналитика базы и финансовая аналитика", bold=True)
    for item in [
        "Регулярно читать базу и воронку: где теряем конверсию, какой сегмент и какой продукт проседают.",
        "Сводить CRM-картину с выгрузками УТ / финансов: план-факт, экономика запуска, мотивация, бюджеты.",
        "На основе цифр принимать решения по приоритетам обзвона, акциям и ресурсным встречам.",
    ]:
        add_bullet(doc, item)

    add_body(doc, "4.3. Продвижение, маркетинг и кросс-функциональные встречи", bold=True)
    for item in [
        "Инициировать и сопровождать мероприятия / акции по продвижению сервисов и маркетинговые стратегии запусков.",
        "Проводить большое число встреч со смежными подразделениями по продуктам портфеля и фиксировать договорённости.",
        "Не допускать, чтобы МПП работали с устаревшим статусом продукта, оффера или ограничений.",
    ]:
        add_bullet(doc, item)

    add_body(doc, "4.4. Управление МПП (обязательный контур внутри продуктовой нагрузки)", bold=True)
    for item in [
        "Держать контроль канбанов и ключевых сделок, особенно по приоритетным продуктам портфеля.",
        "Коротко обучать МПП изменениям продуктов / акций и сохранять стандарт качества.",
        "По возможности увеличивать долю времени на работу с МПП — это ускоряет проверку стратегий на рынке.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. Требования к компетенциям")
    comp = doc.add_table(rows=1, cols=2)
    set_table_full_width(comp)
    write_cell(comp.rows[0].cells[0], "Блок компетенций", size=8.5, bold=True, fill=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(comp.rows[0].cells[1], "Ожидаемый уровень", size=8.5, bold=True, fill=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Управление портфелем 3–5+ продуктов и продуктовые стратегии", "Параллельно ведёт несколько запусков и собирает приоритеты в решения"),
        ("Аналитика базы и воронки", "Видит сегменты, конверсии и узкие места; переводит выводы в действия"),
        ("Финансовая аналитика запусков", "Сводит УТ / план-факт / мотивацию / бюджеты в управленческие решения"),
        ("Маркетинг и мероприятия продвижения сервисов", "Запускает активности и оценивает эффект на базу и сделки"),
        ("Кросс-функциональные встречи по продуктам", "Синхронизирует много стейкхолдеров без потери статуса для команды"),
        ("Управление МПП через канбан / сделки / короткое обучение", "Удерживает исполнение стратегии руками команды на рынке"),
    ]
    for i, (left, right) in enumerate(rows):
        row = comp.add_row()
        bg = WHITE if i % 2 == 0 else ROW_ALT
        write_cell(row.cells[0], left, size=8, fill=bg)
        write_cell(row.cells[1], right, size=8, fill=bg)
    set_col_widths(comp, [12.0, 15.0])

    add_heading(doc, "6. Права")
    for item in [
        "Запрашивать у смежных подразделений данные и доступы, необходимые для управления командой и запусками (включая выгрузки УТ / CRM).",
        "Вносить предложения по изменению продукта, оффера, скриптов, процессов и KPI группы.",
        "Распределять задачи внутри команды МПП и контролировать их исполнение.",
        "Инициировать обучение, наставничество, оценку и кадровые решения в рамках полномочий.",
        "Эскалировать риски по сделкам, продукту, срокам, качеству исполнения и финансовым отклонениям.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Ответственность")
    for item in [
        "За качество управления портфелем продуктов (3–5+) и связанных продуктовых / маркетинговых / финансовых решений.",
        "За качество управления МПП в части канбанов, ключевых сделок и готовности команды к изменениям продуктов.",
        "За качество и своевременность выполнения функций настоящей инструкции.",
        "За соблюдение стандартов продаж, регламентов и корпоративных правил ГК Форус.",
        "За достоверность управленческой и операционной отчётности группы.",
        "За результат обучения и контроля МПП в части стандартов, скриптов и работы с продуктом.",
        "За соблюдение конфиденциальности коммерческой и внутренней информации.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Взаимодействия")
    inter = doc.add_table(rows=1, cols=2)
    set_table_full_width(inter)
    write_cell(inter.rows[0].cells[0], "Контрагент", size=8.5, bold=True, fill=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(inter.rows[0].cells[1], "Предмет взаимодействия", size=8.5, bold=True, fill=HEADER_BG, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, (who, what) in enumerate([
        ("Продуктовая разработка / ЦРП", "Стратегии продуктов портфеля, ограничения, доработки, партнёрский канал"),
        ("Маркетинг", "Мероприятия продвижения, акции, материалы, сегментные кампании, эффект на базу"),
        ("Финансовая аналитика / УТ", "Выгрузки продаж, план-факт, бюджеты, мотивация, экономика запусков"),
        ("Смежные подразделения по продуктам", "Большой поток встреч: статусы, согласования, эскалации по 3–5+ продуктам"),
        ("МПП группы", "Канбаны, ключевые сделки, короткое обучение под изменения продуктов / акций"),
        ("Руководитель направления", "Приоритеты портфеля, KPI, ресурсные решения, эскалации"),
    ]):
        row = inter.add_row()
        bg = WHITE if i % 2 == 0 else ROW_ALT
        write_cell(row.cells[0], who, size=8, bold=True, fill=bg)
        write_cell(row.cells[1], what, size=8, fill=bg)
    set_col_widths(inter, [7.0, 20.0])

    add_heading(doc, "9. Результат работы (ключевые показатели)")
    for item in [
        "Качество и своевременность продуктовых стратегий по портфелю (3–5+ продуктов).",
        "Глубина аналитики базы и влияние выводов на приоритеты обзвона / акций / скриптов.",
        "Качество финансовой аналитики: план-факт, экономика продуктов, мотивация, бюджеты.",
        "Результативность мероприятий и маркетинговых активностей по продвижению сервисов.",
        "Скорость и качество синхронизации со смежными подразделениями по продуктам.",
        "Выполнение плана продаж группы и динамика конверсий по этапам воронки.",
        "Качество канбанов / ключевых сделок МПП и готовность команды к изменениям продуктов.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. Ознакомление")
    sign = doc.add_table(rows=3, cols=2)
    set_table_full_width(sign)
    write_cell(sign.rows[0].cells[0], "Инструкцию утвердил:\n\n________________ / ________________", size=9, fill=WHITE)
    write_cell(sign.rows[0].cells[1], "С инструкцией ознакомлен(а):\n\n________________ / ________________", size=9, fill=WHITE)
    write_cell(sign.rows[1].cells[0], "Должность: _______________________", size=8, color=GRAY, fill=WHITE)
    write_cell(sign.rows[1].cells[1], "Тимлид группы продуктового запуска", size=8, color=GRAY, fill=WHITE)
    write_cell(sign.rows[2].cells[0], "Дата: «____» ______________ 20____ г.", size=8, color=GRAY, fill=WHITE)
    write_cell(sign.rows[2].cells[1], "Дата: «____» ______________ 20____ г.", size=8, color=GRAY, fill=WHITE)
    set_col_widths(sign, [13.5, 13.5])

    foot = doc.add_paragraph()
    clear_paragraph(foot)
    foot.paragraph_format.space_before = Pt(10)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        foot.add_run("© ГК Форус  ·  Внутренний документ  ·  Должностная инструкция тимлида группы продуктового запуска"),
        size=7,
        color=GRAY,
    )

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    ART_RU.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    doc.save(OUT_ROOT)
    doc.save(ART_RU)
    doc.save(ART_EN)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Saved: {path}")
