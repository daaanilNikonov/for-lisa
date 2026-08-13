#!/usr/bin/env python3
"""Build printable B&W team number cards (Word): teams 1–4 × 5 cards."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "квиз 1с-эпд" / "Карточки_команд_1-4.docx"

TEAMS = 4
CARDS_PER_TEAM = 5
COLS = 2
ROWS = 3  # 2×3 grid, 5 cards + 1 empty
CARD_CM = 8.2
GAP_TWIPS = 120  # ~2 mm between cards
PAGE_MARGIN_CM = 1.4


def set_cell_borders(cell, sz: int = 28) -> None:
    """Thick black borders (sz in eighths of a point; 28 ≈ 3.5 pt)."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tc_pr.append(borders)


def clear_cell_borders(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tc_pr.append(borders)


def set_cell_margins(cell, margin_twips: int = 80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(old)
    mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(margin_twips))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def set_vertical_center(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:vAlign")):
        tc_pr.remove(old)
    v = OxmlElement("w:vAlign")
    v.set(qn("w:val"), "center")
    tc_pr.append(v)


def set_row_height(row, cm: float) -> None:
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    for old in tr_pr.findall(qn("w:trHeight")):
        tr_pr.remove(old)
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), "exact")
    tr_pr.append(h)


def fill_card(cell, team: int) -> None:
    set_cell_borders(cell, sz=28)
    set_cell_margins(cell, 80)
    set_vertical_center(cell)

    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("Команда")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(18)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(str(team))
    r2.bold = True
    r2.font.name = "Arial"
    r2.font.size = Pt(108)


def configure_table(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)

    for old in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(old)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    for old in tbl_pr.findall(qn("w:tblCellSpacing")):
        tbl_pr.remove(old)
    spacing = OxmlElement("w:tblCellSpacing")
    spacing.set(qn("w:w"), str(GAP_TWIPS))
    spacing.set(qn("w:type"), "dxa")
    tbl_pr.append(spacing)

    # Keep table centered on page
    for old in tbl_pr.findall(qn("w:jc")):
        tbl_pr.remove(old)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tbl_pr.append(jc)


def add_team_page(doc: Document, team: int, first_page: bool) -> None:
    if not first_page:
        doc.add_page_break()

    table = doc.add_table(rows=ROWS, cols=COLS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    configure_table(table)

    cards = [team] * CARDS_PER_TEAM
    while len(cards) < ROWS * COLS:
        cards.append(None)

    for row_idx in range(ROWS):
        row = table.rows[row_idx]
        set_row_height(row, CARD_CM)
        for col_idx in range(COLS):
            cell = row.cells[col_idx]
            cell.width = Cm(CARD_CM)
            value = cards[row_idx * COLS + col_idx]
            if value is None:
                cell.text = ""
                clear_cell_borders(cell)
                set_vertical_center(cell)
            else:
                fill_card(cell, value)

    # spacer paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def build() -> Path:
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(PAGE_MARGIN_CM)
    section.right_margin = Cm(PAGE_MARGIN_CM)
    section.top_margin = Cm(PAGE_MARGIN_CM)
    section.bottom_margin = Cm(PAGE_MARGIN_CM)

    for i, team in enumerate(range(1, TEAMS + 1)):
        add_team_page(doc, team, first_page=(i == 0))

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Saved: {path} ({path.stat().st_size} bytes)")
    print(f"Cards: {TEAMS} teams × {CARDS_PER_TEAM} = {TEAMS * CARDS_PER_TEAM}")
    print("Layout: 1 team per A4 page, 5 square cards (2×3 grid)")
