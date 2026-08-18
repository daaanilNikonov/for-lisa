#!/usr/bin/env python3
"""Сборка редактируемого DOCX-резюме Попова С. Д."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw

ROOT = Path(__file__).resolve().parents[1]
RESUME_DIR = ROOT / "resume"
PHOTO = RESUME_DIR / "photo.jpg"
CIRCLE = RESUME_DIR / "photo-circle.png"
OUT = RESUME_DIR / "Попов_Семён_Backend_разработчик.docx"

INK = RGBColor(0x15, 0x20, 0x2B)
MUTED = RGBColor(0x5C, 0x67, 0x74)
ACCENT = RGBColor(0x0F, 0x6B, 0x61)
DEEP = RGBColor(0x0C, 0x3D, 0x48)
WHITE = RGBColor(0xF4, 0xF8, 0xF8)
TEAL = RGBColor(0x9E, 0xD9, 0xCF)


def set_run_font(run, name="Calibri", size=11, bold=False, color=INK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_paragraph_spacing(p, before=0, after=4, line=1.15):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, **sides):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for key, val_cm in sides.items():
        node = OxmlElement(f"w:{key}")
        node.set(qn("w:w"), str(int(val_cm * 567)))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_borders(cell, color="FFFFFF", sz="0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil" if sz == "0" else "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def prevent_table_indent(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=4, line=1.0)
    run = p.add_run(text.upper())
    set_run_font(run, size=10, bold=True, color=ACCENT)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E6EBE9")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_job_header(doc, title, org, dates, note=None):
    table = doc.add_table(rows=1, cols=2)
    prevent_table_indent(table)
    left, right = table.rows[0].cells
    set_cell_borders(left)
    set_cell_borders(right)
    set_cell_margins(left, top=0.04, bottom=0.02, left=0, right=0.2)
    set_cell_margins(right, top=0.04, bottom=0.02, left=0.1, right=0)

    p = left.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=1.1)
    run = p.add_run(title)
    set_run_font(run, size=12, bold=True)

    p2 = left.add_paragraph()
    set_paragraph_spacing(p2, before=1, after=0, line=1.1)
    run = p2.add_run(org)
    set_run_font(run, size=11, bold=True, color=DEEP)

    if note:
        p3 = left.add_paragraph()
        set_paragraph_spacing(p3, before=2, after=2, line=1.15)
        run = p3.add_run(note)
        set_run_font(run, size=10, color=MUTED)

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, before=0, after=0, line=1.1)
    run = p.add_run(dates)
    set_run_font(run, size=10, bold=True)
    right.width = Cm(4.4)
    left.width = Cm(13.6)
    return table


def add_body(doc, text, size=10.5, space_after=4, italic=False):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=space_after, line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=size)
    run.italic = italic
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=2, line=1.15)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, size=10.5, bold=True)
        run = p.add_run(text)
        set_run_font(run, size=10.5)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.5)
    return p


def add_case(doc, kicker, title, rows):
    table = doc.add_table(rows=1, cols=1)
    prevent_table_indent(table)
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F3F7F6")
    set_cell_margins(cell, top=0.12, bottom=0.12, left=0.25, right=0.2)
    set_cell_borders(cell, color="0F6B61", sz="12")

    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=1, line=1.05)
    run = p.add_run(kicker.upper())
    set_run_font(run, size=8, bold=True, color=ACCENT)

    p = cell.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4, line=1.1)
    run = p.add_run(title)
    set_run_font(run, size=11, bold=True)

    for label, value in rows:
        p = cell.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.12)
        run = p.add_run(f"{label}: ")
        set_run_font(run, size=10, bold=True, color=ACCENT)
        run = p.add_run(value)
        set_run_font(run, size=10.5)

    # spacer after case
    sp = doc.add_paragraph()
    set_paragraph_spacing(sp, before=0, after=2, line=1.0)
    return table


def make_circle_photo():
    im = Image.open(PHOTO).convert("RGBA")
    size = min(im.size)
    im = im.resize((size, size), Image.Resampling.LANCZOS)
    mask = Image.new("L", (size, size), 0)
    ImageDraw.Draw(mask).ellipse((1, 1, size - 2, size - 2), fill=255)
    out = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    out.paste(im, (0, 0))
    out.putalpha(mask)
    out.save(CIRCLE)
    return CIRCLE


def build():
    make_circle_photo()
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.4)

    # Header with photo
    header = doc.add_table(rows=1, cols=2)
    prevent_table_indent(header)
    text_cell, photo_cell = header.rows[0].cells
    shade_cell(text_cell, "0C3D48")
    shade_cell(photo_cell, "0C3D48")
    set_cell_borders(text_cell)
    set_cell_borders(photo_cell)
    set_cell_margins(text_cell, top=0.28, bottom=0.28, left=0.35, right=0.2)
    set_cell_margins(photo_cell, top=0.22, bottom=0.22, left=0.1, right=0.28)
    text_cell.width = Cm(14.4)
    photo_cell.width = Cm(3.8)

    p = text_cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=2, line=1.05)
    run = p.add_run("Попов Семён Дмитриевич")
    set_run_font(run, size=22, bold=True, color=WHITE)

    p = text_cell.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.05)
    run = p.add_run("Backend-разработчик  ·  Go / Python")
    set_run_font(run, size=12, color=TEAL)

    p = text_cell.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line=1.15)
    run = p.add_run("Иркутск, Россия  ·  +7 (983) 244-61-88  ·  simoneBsn@yandex.ru  ·  t.me/LackOfHapp")
    set_run_font(run, size=9.5, color=RGBColor(0xD5, 0xE4, 0xE6))

    p = photo_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    run = p.add_run()
    run.add_picture(str(CIRCLE), width=Cm(2.55), height=Cm(2.55))

    # Desired position
    add_section_heading(doc, "Желаемая должность")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=1, line=1.1)
    run = p.add_run("Backend-разработчик")
    set_run_font(run, size=14, bold=True)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4, line=1.1)
    run = p.add_run("Программист, разработчик  ·  Полная занятость  ·  Удалённо / гибрид / офис, рассматриваю релокацию")
    set_run_font(run, size=10, color=MUTED)

    # Experience
    add_section_heading(doc, "Опыт работы — 10 месяцев")
    add_job_header(
        doc,
        "Инженер-программист",
        "ООО ЦТО «Торговые технологии»",
        "Декабрь 2025 — настоящее время\n9 месяцев",
        "Внедрение и сопровождение IT-систем в HoReCa. Промышленный контур R-Keeper, iiko, Shelter PMS на объектах заказчиков.",
    )
    add_body(
        doc,
        "В команде внедрения и сопровождения закрываю инциденты на живых объектах: от сети и серверов до рабочего сценария персонала. Задачу не отпускаю, пока клиент не подтвердит результат.",
    )
    add_bullet(doc, "Разворачиваю и сопровождаю специализированное ПО на Windows / Windows Server: кассовые серверы, PMS, складской учёт, торговое оборудование.")
    add_bullet(doc, "Настраиваю сеть объектов (DNS, VPN между точками и серверами) и восстанавливаю сервисы при сбоях.")
    add_bullet(doc, "Диагностирую отказы R-Keeper, iiko, Shelter PMS по логам и конфигурации: отделяю дефект от штатного поведения, сверяюсь с документацией вендора, при необходимости привлекаю профильных коллег.")
    add_bullet(doc, "Консультирую ресепшен, бухгалтерию и склад — перевожу логику системы в понятный рабочий сценарий и проверяю его вместе с клиентом.")

    add_case(
        doc,
        "Кейс · отель",
        "Заселение гостя по цифровому ID Max — без ручного ввода паспортных данных",
        [
            ("Команда", "Внедрял фичу в действующий контур вместе со стойкой: обновление PMS, оборудование, обучение администраторов."),
            ("Личный вклад", "Обновил систему, подключил и настроил сканер QR, проверил цепочку «сканирование → карточка гостя» и закрепил сценарий у персонала."),
            ("Результат", "Стойка перестала вбивать ФИО и паспорт руками. Меньше ошибок в карточке гостя, быстрее заселение, короче очередь в часы заезда."),
        ],
    )
    add_case(
        doc,
        "Кейс · склад",
        "Искажение складского учёта в актах переработки",
        [
            ("Команда", "Разобрал кейс с бухгалтерией склада: документы, номенклатура и логика акта переработки в WMS."),
            ("Личный вклад", "Нашёл, почему 2 кг курицы «превращались» в 2 кг филе без потерь, и восстановил корректный расчёт выхода продукции."),
            ("Результат", "Акты снова отражают реальный выход, а не копируют вес сырья. Честные остатки, себестоимость и отчётность — без завышенного объёма готовой продукции."),
        ],
    )

    add_job_header(doc, "Стажёр", "ГК «Прогресс»", "Июль 2024 — август 2024\n1 месяц")
    add_bullet(doc, "Диагностика и обслуживание офисной инфраструктуры; участие в настройке сети на MikroTik (DHCP, NAT, Firewall, VPN).")
    add_bullet(doc, " закрывал типовые инциденты по оборудованию и сети, не останавливая работу офиса.", bold_prefix="Личный результат:")

    # Projects
    add_section_heading(doc, "Проекты")
    add_job_header(doc, "Компоненты 2D-игрового движка на C++ / SDL3", "Дипломный проект · НИ ТПУ", "2025")
    add_body(doc, "Команда: делили проект по направлениям: на своём был ведущим, на смежном поддерживал команду. Согласовывал требования и дедлайны, собирал концепцию и доводил продукт до защиты.")
    add_body(doc, "Личный вклад: формализовал сущности и связи, спроектировал логику компонентов, реализовал их на C++ / SDL3, участвовал в тестировании, сборке и защите.")
    add_body(doc, "Результат: защищённый диплом, рабочий набор компонентов 2D-движка. Средний балл обучения — 4.46.")

    add_job_header(doc, "Backend-практика: REST API на Go", "Личные и учебные проекты", "2025 — н.в.")
    add_body(doc, "Постановка: до кода прорабатываю сущности, связи, потоки данных и контракты API. Отсекаю избыточные решения в пользу простых и устойчивых.")
    add_body(doc, "Личный вклад: пишу обработчики и бизнес-логику на Go, проектирую REST-эндпоинты, решаю задачи на SQL (выборки, агрегации, JOIN). Автоматизирую Fedora скриптами Bash / Zsh, осваиваю Docker.")
    add_body(doc, "Результат: умение провести задачу от модели предметной области до работающего API и сверить поведение кода с ожидаемой логикой.")

    # Education
    add_section_heading(doc, "Образование")
    add_job_header(
        doc,
        "Магистратура · Корпоративные информационные системы. Инновационные методики и платформы",
        "ИРНИТУ",
        "2025 — настоящее время",
    )
    add_job_header(
        doc,
        "Бакалавр · Информатика и вычислительная техника",
        "ФГАОУ ВО НИ Томский политехнический университет",
        "2021 — 2025",
        "Профиль «Информационно-коммуникационные системы». Средний балл — 4.46. Диплом: «Разработка компонентов 2D-игрового движка на C++ / SDL3».",
    )

    add_section_heading(doc, "Повышение квалификации, курсы")
    add_job_header(
        doc,
        "Профессиональная переподготовка · Разработка методов вычислительного интеллекта на Python",
        "ФГАОУ ВО НИ Томский политехнический университет · диплом о переподготовке",
        "",
    )

    add_section_heading(doc, "Ключевые навыки")
    skills = [
        ("Backend", "Go, Python, REST API, SQL, PostgreSQL, Git, Docker"),
        ("Инженерия", "Linux, Bash / Zsh, Windows Server, DNS, VPN, MikroTik, Postman"),
        ("Разработка", "C++, SDL3, алгоритмы и структуры данных"),
        ("Домен", "R-Keeper, iiko, Shelter PMS, кассовое и торговое оборудование"),
    ]
    for label, value in skills:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.15)
        run = p.add_run(f"{label}:  ")
        set_run_font(run, size=10, bold=True, color=ACCENT)
        run = p.add_run(value)
        set_run_font(run, size=10.5)

    add_section_heading(doc, "О себе")
    add_body(
        doc,
        "Бакалавр НИ ТПУ, «Информатика и вычислительная техника», профиль «Информационно-коммуникационные системы», средний балл — 4.46. "
        "Диплом: «Разработка компонентов 2D-игрового движка на C++ / SDL3». "
        "Переподготовка НИ ТПУ: «Разработка методов вычислительного интеллекта на Python». "
        "Магистратура ИРНИТУ: «Корпоративные информационные системы. Инновационные методики и платформы». "
        "Стек: Python, Go, PostgreSQL, Docker, Git, Linux, Windows Server. Английский язык — B2.",
    )

    add_section_heading(doc, "Знание языков")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line=1.15)
    run = p.add_run("Русский")
    set_run_font(run, size=11, bold=True)
    run = p.add_run(" — родной     ")
    set_run_font(run, size=11)
    run = p.add_run("Английский")
    set_run_font(run, size=11, bold=True)
    run = p.add_run(" — B2")
    set_run_font(run, size=11)

    doc.save(OUT)
    print(f"Written {OUT}")


if __name__ == "__main__":
    build()
