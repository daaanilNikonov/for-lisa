#!/usr/bin/env python3
"""Коммерческое предложение-листовка: 1С:Кабинет сотрудника (750).

PDF в стиле листовок Форус + DOCX в том же содержании.
Менеджер: Данил Кургузов.
5 часов линии в подарок: ~4 на настройку, 1 на доп. вопросы. Клиент за часы не платит.
"""

from __future__ import annotations

from pathlib import Path

from reportlab.lib.colors import HexColor, white, Color
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor as DocRGB

ROOT = Path(__file__).resolve().parents[1]
OUT_PDF = ROOT / "output" / "КП_ТФМ_Спецтехника_КЭДО_750.pdf"
OUT_DOCX = ROOT / "output" / "КП_ТФМ_Спецтехника_КЭДО_750.docx"
BRAND = ROOT / "assets_forus" / "brand"

YELLOW = HexColor("#F0C14A")
YELLOW_SOFT = HexColor("#FFF8E8")
DARK = HexColor("#1A1A1A")
GRAY = HexColor("#5A5A5A")
LIGHT = HexColor("#F7F7F7")
LINE = HexColor("#E5E5E5")
GREEN_BG = HexColor("#EAF6EA")
GREEN = HexColor("#1F6B2E")
RED_BG = HexColor("#F9EAEA")
RED = HexColor("#8C2020")

pdfmetrics.registerFont(TTFont("DejaVu", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"))
pdfmetrics.registerFont(TTFont("DejaVuBold", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"))

PAGE_W, PAGE_H = A4
ML, MR = 14 * mm, 14 * mm
MT, MB = 12 * mm, 12 * mm
CW = PAGE_W - ML - MR

CABINETS = 223_200
GIFT_HOURS = 5
GIFT_VALUE = 18_300
TOTAL = CABINETS

MANAGER = "Данил Кургузов"
MANAGER_EMAIL = "dkurguzov@forus.ru"
MANAGER_PHONE = "+7 (3952) 78-00-00"


# ---------- helpers PDF ----------

def wave(c, x, y, w, h, flip=False):
    c.setFillColor(YELLOW)
    p = c.beginPath()
    if not flip:
        p.moveTo(x, y)
        p.curveTo(x + w * 0.3, y + h, x + w * 0.55, y - h * 0.2, x + w * 0.8, y + h * 0.7)
        p.curveTo(x + w * 0.92, y + h, x + w, y + h * 0.3, x + w, y + h * 0.5)
        p.lineTo(x + w, y + h * 1.5)
        p.lineTo(x, y + h * 1.5)
        p.close()
    else:
        p.moveTo(x, y + h)
        p.curveTo(x + w * 0.25, y, x + w * 0.5, y + h * 1.1, x + w * 0.75, y + h * 0.25)
        p.curveTo(x + w * 0.9, y, x + w, y + h * 0.4, x + w, y)
        p.lineTo(x + w, y - h)
        p.lineTo(x, y - h)
        p.close()
    c.drawPath(p, fill=1, stroke=0)


def wrap(c, text, font, size, max_w):
    words, lines, cur = text.split(), [], ""
    for w in words:
        t = f"{cur} {w}".strip()
        if c.stringWidth(t, font, size) <= max_w:
            cur = t
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines


class Layout:
    def __init__(self, c: canvas.Canvas):
        self.c = c
        self.y = PAGE_H - MT

    def gap(self, h):
        self.y -= h

    def ensure(self, need):
        """If not enough space, caller should have started new page."""
        return self.y - MB >= need

    def rule(self, thick=2.2):
        self.c.setFillColor(YELLOW)
        self.c.rect(ML, self.y, CW, thick, fill=1, stroke=0)
        self.y -= 4 * mm

    def title(self, text, size=11):
        self.c.setFont("DejaVuBold", size)
        self.c.setFillColor(DARK)
        self.c.drawString(ML, self.y, text)
        self.y -= 2.2 * mm
        self.rule(1.6)

    def text(self, text, size=8.2, color=DARK, leading=None, bold=False, indent=0):
        leading = leading or (size + 2.4)
        font = "DejaVuBold" if bold else "DejaVu"
        self.c.setFont(font, size)
        self.c.setFillColor(color)
        for line in wrap(self.c, text, font, size, CW - indent):
            self.c.drawString(ML + indent, self.y, line)
            self.y -= leading * mm / 2.834  # wrong - leading should be in points
        # fix: use points properly
        return

    def para(self, text, size=8.2, color=DARK, leading=11, bold=False, indent=0):
        font = "DejaVuBold" if bold else "DejaVu"
        self.c.setFont(font, size)
        self.c.setFillColor(color)
        for line in wrap(self.c, text, font, size, CW - indent):
            self.c.drawString(ML + indent, self.y, line)
            self.y -= leading

    def check(self, text, size=8, leading=11):
        self.c.setFillColor(YELLOW)
        self.c.setFont("DejaVuBold", size + 1)
        self.c.drawString(ML, self.y - 1, "✓")
        font = "DejaVu"
        self.c.setFont(font, size)
        self.c.setFillColor(DARK)
        x0 = ML + 5 * mm
        lines = wrap(self.c, text, font, size, CW - 5 * mm)
        for i, line in enumerate(lines):
            self.c.drawString(x0, self.y, line)
            self.y -= leading
        self.y -= 1.5

    def box(self, h, fill=YELLOW_SOFT, stroke=YELLOW, pad=False):
        self.c.setFillColor(fill)
        self.c.setStrokeColor(stroke)
        self.c.setLineWidth(1.2)
        self.c.roundRect(ML, self.y - h, CW, h, 5, fill=1, stroke=1)
        return self.y - h


def header(c):
    logo = BRAND / "forus_logo_word.png"
    if logo.exists():
        c.drawImage(str(logo), ML, PAGE_H - MT - 12 * mm, width=34 * mm, height=12 * mm,
                    mask="auto", preserveAspectRatio=True, anchor="sw")
    c.setFillColor(GRAY)
    c.setFont("DejaVu", 7.2)
    lines = [
        "Группа компаний «Форус»",
        "Центр компетенции по кадровому электронному документообороту",
        f"{MANAGER_PHONE}   ·   www.forus.ru",
    ]
    yy = PAGE_H - MT - 3 * mm
    for line in lines:
        c.drawRightString(PAGE_W - MR, yy, line)
        yy -= 9
    wave(c, PAGE_W - 52 * mm, PAGE_H - 14 * mm, 52 * mm, 7 * mm)
    c.setFillColor(YELLOW)
    c.rect(ML, PAGE_H - MT - 14.5 * mm, CW, 2.4, fill=1, stroke=0)


def footer(c, page, pages=2):
    wave(c, 0, 2 * mm, 48 * mm, 6 * mm, flip=True)
    c.setFillColor(YELLOW)
    c.rect(ML, MB - 1 * mm, CW, 1.8, fill=1, stroke=0)
    c.setFont("DejaVu", 6.8)
    c.setFillColor(GRAY)
    c.drawString(ML, MB - 5.5 * mm, "www.forus.ru  ·  г. Иркутск, ул. Ямская, 1/1")
    c.drawRightString(PAGE_W - MR, MB - 5.5 * mm, f"{page} / {pages}")


def draw_benefit_card(c, x, y, w, h, title, body):
    c.setFillColor(LIGHT)
    c.roundRect(x, y - h, w, h, 4, fill=1, stroke=0)
    c.setFillColor(YELLOW)
    c.rect(x, y - h, 2 * mm, h, fill=1, stroke=0)
    c.setFillColor(DARK)
    c.setFont("DejaVuBold", 7.8)
    c.drawString(x + 4 * mm, y - 4.5 * mm, title)
    c.setFont("DejaVu", 6.9)
    c.setFillColor(GRAY)
    yy = y - 8.5 * mm
    for line in wrap(c, body, "DejaVu", 6.9, w - 7 * mm):
        c.drawString(x + 4 * mm, yy, line)
        yy -= 9


def build_pdf():
    OUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(OUT_PDF), pagesize=A4)
    c.setTitle("Коммерческое предложение — 1С:Кабинет сотрудника")
    c.setAuthor(f"ГК Форус · {MANAGER}")

    # ===== PAGE 1 =====
    header(c)
    L = Layout(c)
    L.y = PAGE_H - MT - 18 * mm

    c.setFont("DejaVuBold", 18)
    c.setFillColor(DARK)
    c.drawCentredString(PAGE_W / 2, L.y, "1С:Кабинет сотрудника")
    L.gap(6 * mm)
    c.setFont("DejaVu", 9.5)
    c.setFillColor(GRAY)
    c.drawCentredString(PAGE_W / 2, L.y, "Коммерческое предложение")
    L.gap(5 * mm)
    c.setFont("DejaVuBold", 9)
    c.setFillColor(DARK)
    c.drawCentredString(
        PAGE_W / 2, L.y,
        "Автоматизируйте рутину — экономьте до 70% времени и сократите траты на расходники",
    )
    L.gap(7 * mm)

    # hero with image if available
    devices = BRAND / "devices_mockup.png"
    if not devices.exists():
        devices = BRAND / "leaflet_img_1.png"
    box_h = 28 * mm
    c.setFillColor(YELLOW_SOFT)
    c.setStrokeColor(YELLOW)
    c.setLineWidth(1.3)
    c.roundRect(ML, L.y - box_h, CW, box_h, 6, fill=1, stroke=1)
    text_w = CW - 48 * mm if devices.exists() else CW - 6 * mm
    c.setFillColor(DARK)
    c.setFont("DejaVuBold", 9)
    c.drawString(ML + 4 * mm, L.y - 5 * mm, "Почему компании переходят на кадровый электронный документооборот")
    yy = L.y - 10 * mm
    for t in [
        "Согласования занимают минуты, а не дни.",
        "Меньше бумаги, печати, курьеров и риска потерять оригинал.",
        "Удалённые сотрудники подписывают документы без поездок в офис.",
        "Кадровики и бухгалтерия тратят меньше времени на рутину.",
    ]:
        c.setFillColor(YELLOW)
        c.setFont("DejaVuBold", 8)
        c.drawString(ML + 4 * mm, yy, "✓")
        c.setFillColor(DARK)
        c.setFont("DejaVu", 7.6)
        c.drawString(ML + 8 * mm, yy, t)
        yy -= 4.2 * mm
    if devices.exists():
        c.drawImage(
            str(devices), PAGE_W - MR - 44 * mm, L.y - box_h + 2 * mm,
            width=42 * mm, height=24 * mm, mask="auto", preserveAspectRatio=True, anchor="sw",
        )
    L.y -= box_h + 5 * mm

    L.title("Что умеет сервис")
    benefits = [
        ("Удобная работа с «неуловимыми» сотрудниками",
         "Расчётные листки, справки и запросы по отпускам доступны удалённо каждому сотруднику в личном кабинете."),
        ("Ознакомление с документами в один клик",
         "Не нужно распечатывать, собирать подписи и хранить горы бумаг. Отправили документ — сотрудник ознакомился и подтвердил получение."),
        ("Удалённый приём на работу",
         "Новый сотрудник оформляет трудовой договор, заявление и ознакомление с правилами дистанционно, не приезжая в офис."),
        ("Меньше ручного ввода и ошибок",
         "Данные от сотрудников автоматически попадают в 1С. Заявления превращаются в готовые приказы и записи."),
        ("Согласование без походов по кабинетам",
         "Отпуск, командировка, отгул — руководитель согласовывает с телефона, данные сразу уходят в 1С."),
        ("Общение без сторонних мессенджеров",
         "Рабочие вопросы и документы — внутри системы, с соблюдением требований закона о персональных данных."),
        ("Электронный архив и защита от штрафов",
         "Кадровые документы хранятся в электронном виде. Есть подтверждение, что сотрудник получил расчётный листок и ознакомился с приказом."),
        ("Простое внедрение для ИТ-отдела",
         "Сервис встроен в 1С: не нужна отдельная кадровая платформа, новые клиентские лицензии 1С и сложная интеграция «с нуля»."),
    ]
    card_w = (CW - 3 * mm) / 2
    card_h = 20 * mm
    for i in range(0, 8, 2):
        for col in range(2):
            title, body = benefits[i + col]
            x = ML + col * (card_w + 3 * mm)
            draw_benefit_card(c, x, L.y, card_w, card_h, title, body)
        L.y -= card_h + 2.5 * mm

    L.gap(2 * mm)
    L.title("Акция «Больше, чем кешбэк!»")
    promo_h = 26 * mm
    c.setFillColor(DARK)
    c.roundRect(ML, L.y - promo_h, CW, promo_h, 6, fill=1, stroke=0)
    c.setFillColor(YELLOW)
    c.setFont("DejaVuBold", 9.5)
    c.drawString(ML + 4 * mm, L.y - 5 * mm, "5 часов линии консультаций — в подарок")
    c.setFillColor(white)
    c.setFont("DejaVu", 7.8)
    lines = [
        "При оплате пакета личных кабинетов мы начисляем на баланс линии консультаций 5 часов бесплатно",
        f"(выгода {GIFT_VALUE:,} ₽).".replace(",", " "),
        "Из них в среднем 4 часа уходят на настройку и запуск сервиса.",
        "Ещё 1 час остаётся у вас на вопросы, консультации и решение возникающих задач по 1С",
        "с нашими специалистами.",
        "Отдельно за часы линии консультаций платить не нужно — они полностью покрываются подарком.",
    ]
    yy = L.y - 9.5 * mm
    for line in lines:
        c.drawString(ML + 4 * mm, yy, line)
        yy -= 3.3 * mm
    L.y -= promo_h + 5 * mm

    L.title("Ваш пакет: 750 личных кабинетов на 12 месяцев")
    c.setFont("DejaVu", 7.8)
    c.setFillColor(GRAY)
    c.drawString(ML, L.y, "Состав лицензий: 500 + 200 + 50 кабинетов. Цена за сотрудника — около 25 ₽ в месяц.")
    L.gap(4 * mm)

    # price strip
    c.setFillColor(YELLOW)
    c.roundRect(ML, L.y - 14 * mm, CW, 14 * mm, 5, fill=1, stroke=0)
    c.setFillColor(DARK)
    c.setFont("DejaVuBold", 10)
    c.drawString(ML + 4 * mm, L.y - 5.5 * mm, "1С:Кабинет сотрудника — 750 кабинетов / 12 месяцев")
    c.setFont("DejaVuBold", 14)
    c.drawRightString(PAGE_W - MR - 4 * mm, L.y - 6 * mm, f"{TOTAL:,} ₽".replace(",", " "))
    c.setFont("DejaVu", 7.5)
    c.drawString(ML + 4 * mm, L.y - 11 * mm, "Настройка, запуск и час поддержки по вопросам 1С — бесплатно в рамках подарочных часов")
    L.y -= 18 * mm

    footer(c, 1)

    # ===== PAGE 2 =====
    c.showPage()
    header(c)
    L = Layout(c)
    L.y = PAGE_H - MT - 18 * mm

    L.title("1. Условия")
    for t in [
        "Подключаем сервис «1С:Кабинет сотрудника» на 750 личных кабинетов сроком на 12 месяцев.",
        "В запуск входит: подключение к вашей 1С, настройка ролей и правил, помощь с выпуском электронных подписей, "
        "запуск пилотной группы и инструкции для сотрудников.",
        "5 часов линии консультаций начисляются в подарок. Около 4 часов — настройка и запуск; 1 час остаётся "
        "на дополнительные вопросы и помощь по 1С.",
        "За часы линии консультаций клиент не платит.",
        "Электронная подпись для сотрудников (усиленная неквалифицированная) — бесплатно.",
        "Предоставляем шаблоны документов для перехода: положение, уведомления, согласия.",
        "Работы выполняются удалённо. Выезд и нетиповые доработки печатных форм — по отдельному согласованию.",
        "Коммерческое предложение действует 30 календарных дней с даты направления.",
    ]:
        L.check(t, size=7.4, leading=9.6)
    L.gap(3 * mm)

    L.title("2. Сроки")
    c.setFont("DejaVu", 7.5)
    c.setFillColor(GRAY)
    c.drawString(ML, L.y, "Внедрение простое: без отдельной платформы и без долгого проектного цикла.")
    L.gap(4 * mm)

    rows = [
        ("Договор и доступы к 1С", "2–4 рабочих дня", "Можно начинать настройку"),
        ("Настройка и запуск (около 4 часов)", "1 рабочий день", "Сервис подключен, роли и подписи настроены"),
        ("Пилот на небольшой группе", "3–7 рабочих дней", "Проверен рабочий сценарий"),
        ("Обучение ключевых сотрудников", "параллельно", "Видеоуроки и короткие инструкции"),
        ("Подключение всех 750 сотрудников", "2–3 недели", "Массовый запуск личных кабинетов"),
        ("Резерв 1 час поддержки", "по запросу", "Вопросы и помощь по 1С после запуска"),
    ]
    cols = [58 * mm, 40 * mm, CW - 98 * mm]
    # header
    hh = 7 * mm
    c.setFillColor(DARK)
    c.rect(ML, L.y - hh, CW, hh, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont("DejaVuBold", 7.5)
    xs = [ML, ML + cols[0], ML + cols[0] + cols[1]]
    for i, h in enumerate(["Этап", "Срок", "Результат"]):
        c.drawString(xs[i] + 2 * mm, L.y - hh + 2.2 * mm, h)
    L.y -= hh
    for ri, (a, b, d) in enumerate(rows):
        rh = 6.8 * mm
        c.setFillColor(YELLOW_SOFT if ri in (1, 5) else (LIGHT if ri % 2 == 0 else white))
        c.rect(ML, L.y - rh, CW, rh, fill=1, stroke=0)
        c.setStrokeColor(LINE)
        c.setLineWidth(0.4)
        c.rect(ML, L.y - rh, CW, rh, fill=0, stroke=1)
        c.setFillColor(DARK)
        c.setFont("DejaVuBold", 7)
        c.drawString(xs[0] + 2 * mm, L.y - rh + 2.5 * mm, a)
        c.setFillColor(HexColor("#8A6A12"))
        c.setFont("DejaVuBold", 7)
        c.drawString(xs[1] + 2 * mm, L.y - rh + 2.5 * mm, b)
        c.setFillColor(GRAY)
        c.setFont("DejaVu", 6.8)
        c.drawString(xs[2] + 2 * mm, L.y - rh + 2.5 * mm, d)
        L.y -= rh

    L.gap(4 * mm)
    c.setFillColor(GREEN_BG)
    c.setStrokeColor(GREEN)
    c.setLineWidth(1)
    c.roundRect(ML, L.y - 10 * mm, CW, 10 * mm, 4, fill=1, stroke=1)
    c.setFillColor(GREEN)
    c.setFont("DejaVuBold", 7.8)
    c.drawString(ML + 3 * mm, L.y - 4 * mm, "Итоговый ориентир")
    c.setFont("DejaVu", 7.5)
    c.drawString(ML + 35 * mm, L.y - 4 * mm, "настройка за 1 день · пилот за 1–2 недели · все 750 кабинетов за 1–1,5 месяца")
    c.setFont("DejaVu", 6.8)
    c.drawString(ML + 3 * mm, L.y - 7.8 * mm, "Не нужно выделять большую проектную команду и внедрять отдельную кадровую систему.")
    L.y -= 14 * mm

    L.title("3. Бюджет")
    # table
    headers = ["Статья", "Пояснение", "Сумма"]
    data = [
        ("Личные кабинеты, 750 шт. / 12 месяцев", "Пакеты 500 + 200 + 50", f"{TOTAL:,} ₽".replace(",", " ")),
        ("Настройка и запуск (~4 часа)", "Из подарочных часов линии консультаций", "0 ₽"),
        ("Резерв на вопросы по 1С (1 час)", "Из подарочных часов линии консультаций", "0 ₽"),
        ("Подарок: 5 часов линии консультаций", f"Акция «Больше, чем кешбэк!» · выгода {GIFT_VALUE:,} ₽".replace(",", " "), "0 ₽"),
    ]
    cols = [68 * mm, 78 * mm, CW - 146 * mm]
    hh = 7 * mm
    c.setFillColor(DARK)
    c.rect(ML, L.y - hh, CW, hh, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont("DejaVuBold", 7.5)
    xs = [ML, ML + cols[0], ML + cols[0] + cols[1]]
    for i, h in enumerate(headers):
        c.drawString(xs[i] + 2 * mm, L.y - hh + 2.2 * mm, h)
    L.y -= hh
    for ri, (a, b, d) in enumerate(data):
        rh = 6.8 * mm
        c.setFillColor(YELLOW_SOFT if ri > 0 else LIGHT)
        c.rect(ML, L.y - rh, CW, rh, fill=1, stroke=0)
        c.setStrokeColor(LINE)
        c.setLineWidth(0.4)
        c.rect(ML, L.y - rh, CW, rh, fill=0, stroke=1)
        c.setFillColor(DARK)
        c.setFont("DejaVuBold", 7)
        c.drawString(xs[0] + 2 * mm, L.y - rh + 2.7 * mm, a[:42])
        c.setFont("DejaVu", 6.7)
        c.setFillColor(GRAY)
        c.drawString(xs[1] + 2 * mm, L.y - rh + 2.7 * mm, b[:48])
        c.setFont("DejaVuBold", 8)
        c.setFillColor(GREEN if ri > 0 else DARK)
        c.drawRightString(PAGE_W - MR - 2 * mm, L.y - rh + 2.7 * mm, d)
        L.y -= rh

    L.gap(4 * mm)
    c.setFillColor(DARK)
    c.roundRect(ML, L.y - 16 * mm, CW, 16 * mm, 6, fill=1, stroke=0)
    c.setFillColor(YELLOW)
    c.setFont("DejaVuBold", 10)
    c.drawString(ML + 4 * mm, L.y - 6 * mm, "Итого к оплате")
    c.setFont("DejaVuBold", 16)
    c.drawRightString(PAGE_W - MR - 4 * mm, L.y - 6.5 * mm, f"{TOTAL:,} ₽".replace(",", " "))
    c.setFillColor(white)
    c.setFont("DejaVu", 7.5)
    c.drawString(ML + 4 * mm, L.y - 12 * mm, "Только пакет кабинетов. Часы линии консультаций — бесплатно. Подпись сотрудникам — бесплатно.")
    L.y -= 20 * mm

    L.title("Сравнение с другими решениями")
    cmp = [
        ("Сложность запуска", "Типовая настройка около 4 часов", "Часто долгое внедрение новой системы"),
        ("Где работают сотрудники", "В привычной 1С и личном кабинете", "Отдельная кадровая платформа"),
        ("Нагрузка на ИТ-отдел", "Без новых клиентских лицензий 1С", "Новые доступы, интеграции, сопровождение"),
        ("Где хранятся документы", "В вашей базе 1С", "Облако оператора, часто за отдельную плату"),
        ("Оплата запуска", "Часы поддержки — в подарок", "Внедрение обычно оплачивается отдельно"),
    ]
    cols = [42 * mm, 72 * mm, CW - 114 * mm]
    hh = 6.5 * mm
    c.setFillColor(DARK)
    c.rect(ML, L.y - hh, CW, hh, fill=1, stroke=0)
    c.setFillColor(YELLOW)
    c.rect(ML + cols[0], L.y - hh, cols[1], hh, fill=1, stroke=0)
    c.setFont("DejaVuBold", 7)
    xs = [ML, ML + cols[0], ML + cols[0] + cols[1]]
    for i, h in enumerate(["Критерий", "Форус + 1С:Кабинет сотрудника", "Другие операторы"]):
        c.setFillColor(DARK if i == 1 else white)
        c.drawCentredString(xs[i] + cols[i] / 2, L.y - hh + 2 * mm, h)
    L.y -= hh
    for ri, (a, b, d) in enumerate(cmp):
        rh = 6.2 * mm
        for ci, (cell, w) in enumerate(zip((a, b, d), cols)):
            x = xs[ci]
            if ci == 1:
                bg, tc = GREEN_BG, GREEN
            elif ci == 2:
                bg, tc = RED_BG, RED
            else:
                bg, tc = (LIGHT if ri % 2 == 0 else white), DARK
            c.setFillColor(bg)
            c.rect(x, L.y - rh, w, rh, fill=1, stroke=0)
            c.setStrokeColor(LINE)
            c.setLineWidth(0.3)
            c.rect(x, L.y - rh, w, rh, fill=0, stroke=1)
            c.setFillColor(tc)
            c.setFont("DejaVuBold" if ci == 0 else "DejaVu", 6.4)
            if ci == 0:
                c.drawString(x + 1.5 * mm, L.y - rh + 2.3 * mm, cell)
            else:
                c.drawCentredString(x + w / 2, L.y - rh + 2.3 * mm, cell)
        L.y -= rh

    # CTA + manager — фиксируем у низа страницы, чтобы блок не «уплывал»
    cta_top = MB + 34 * mm
    c.setFillColor(DARK)
    c.roundRect(ML, cta_top - 28 * mm, CW, 28 * mm, 6, fill=1, stroke=0)
    c.setFillColor(YELLOW)
    c.setFont("DejaVuBold", 10)
    c.drawCentredString(PAGE_W / 2, cta_top - 6 * mm, "Готовы подключить 750 кабинетов и помочь с запуском")
    c.setFillColor(white)
    c.setFont("DejaVu", 7.5)
    c.drawCentredString(PAGE_W / 2, cta_top - 11 * mm, "Проведём демонстрацию для кадровой службы и ИТ-отдела и ответим на вопросы.")

    c.setFillColor(YELLOW_SOFT)
    c.roundRect(ML + 4 * mm, cta_top - 25 * mm, CW - 8 * mm, 12 * mm, 4, fill=1, stroke=0)
    c.setFillColor(DARK)
    c.setFont("DejaVuBold", 8.5)
    c.drawString(ML + 6 * mm, cta_top - 17 * mm, f"Ваш менеджер: {MANAGER}")
    c.setFont("DejaVu", 7.3)
    c.setFillColor(GRAY)
    c.drawString(ML + 6 * mm, cta_top - 21.5 * mm, f"{MANAGER_EMAIL}   ·   {MANAGER_PHONE}")
    c.setFillColor(DARK)
    c.setFont("DejaVuBold", 7.5)
    c.drawRightString(PAGE_W - MR - 6 * mm, cta_top - 17 * mm, "ГК «Форус»")
    c.setFont("DejaVu", 7)
    c.setFillColor(GRAY)
    c.drawRightString(PAGE_W - MR - 6 * mm, cta_top - 21.5 * mm, "www.forus.ru")

    footer(c, 2)
    c.save()
    print("PDF:", OUT_PDF)


# ---------- DOCX ----------

def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def borders(table, color="CCCCCC"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders_el = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders_el.append(el)
    tblPr.append(borders_el)


def run_font(run, size=10, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_p(doc, text, *, size=10, bold=False, color=DocRGB(0x1A, 0x1A, 0x1A), space_after=6, align="left"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    run_font(r, size=size, bold=bold, color=color)
    return p


def add_check(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    m = p.add_run("✓  ")
    run_font(m, size=10, bold=True, color=DocRGB(0xF0, 0xC1, 0x4A))
    r = p.add_run(text)
    run_font(r, size=10, color=DocRGB(0x1A, 0x1A, 0x1A))


def yellow_underline(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "F0C14A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(14)
    sec.right_margin = Mm(14)
    sec.top_margin = Mm(12)
    sec.bottom_margin = Mm(12)

    # header table
    ht = doc.add_table(rows=1, cols=2)
    left, right = ht.rows[0].cells
    left.paragraphs[0].clear()
    logo = BRAND / "forus_logo_word.png"
    if logo.exists():
        left.paragraphs[0].add_run().add_picture(str(logo), width=Cm(3.6))
    else:
        r = left.paragraphs[0].add_run("Форус")
        run_font(r, size=18, bold=True)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, line in enumerate([
        "Группа компаний «Форус»",
        "Центр компетенции по кадровому электронному документообороту",
        f"{MANAGER_PHONE}  ·  www.forus.ru",
    ]):
        p = right.paragraphs[0] if i == 0 else right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        run_font(r, size=8, bold=(i == 0), color=DocRGB(0x5A, 0x5A, 0x5A) if i else DocRGB(0x1A, 0x1A, 0x1A))

    rule = doc.add_paragraph()
    yellow_underline(rule)
    rule.paragraph_format.space_after = Pt(8)

    add_p(doc, "1С:Кабинет сотрудника", size=18, bold=True, align="center", space_after=2)
    add_p(doc, "Коммерческое предложение", size=11, color=DocRGB(0x5A, 0x5A, 0x5A), align="center", space_after=4)
    add_p(
        doc,
        "Автоматизируйте рутину — экономьте до 70% времени и сократите траты на расходники",
        size=10, bold=True, align="center", space_after=10,
    )

    add_p(doc, "Почему компании переходят на кадровый электронный документооборот", size=12, bold=True, space_after=4)
    yellow_underline(doc.paragraphs[-1])
    for t in [
        "Согласования занимают минуты, а не дни.",
        "Меньше бумаги, печати, курьеров и риска потерять оригинал.",
        "Удалённые сотрудники подписывают документы без поездок в офис.",
        "Кадровики и бухгалтерия тратят меньше времени на рутину.",
    ]:
        add_check(doc, t)

    add_p(doc, "Что умеет сервис", size=12, bold=True, space_after=4, color=DocRGB(0x1A, 0x1A, 0x1A))
    yellow_underline(doc.paragraphs[-1])
    features = [
        ("Удобная работа с «неуловимыми» сотрудниками",
         "Расчётные листки, справки и запросы по отпускам доступны удалённо каждому сотруднику в личном кабинете."),
        ("Ознакомление с документами в один клик",
         "Не нужно распечатывать, собирать подписи и хранить горы бумаг. Отправили документ — сотрудник ознакомился и подтвердил получение."),
        ("Удалённый приём на работу",
         "Новый сотрудник оформляет трудовой договор, заявление и ознакомление с правилами дистанционно."),
        ("Меньше ручного ввода и ошибок",
         "Данные от сотрудников автоматически попадают в 1С. Заявления превращаются в готовые приказы и записи."),
        ("Согласование без походов по кабинетам",
         "Отпуск, командировка, отгул — руководитель согласовывает с телефона, данные сразу уходят в 1С."),
        ("Общение без сторонних мессенджеров",
         "Рабочие вопросы и документы — внутри системы, с соблюдением требований закона о персональных данных."),
        ("Электронный архив и защита от штрафов",
         "Кадровые документы хранятся в электронном виде. Есть подтверждение получения расчётного листка и ознакомления с приказом."),
        ("Простое внедрение для ИТ-отдела",
         "Сервис встроен в 1С: не нужна отдельная кадровая платформа, новые клиентские лицензии 1С и сложная интеграция."),
    ]
    for title, body in features:
        add_p(doc, title, size=10, bold=True, space_after=1)
        add_p(doc, body, size=9, color=DocRGB(0x5A, 0x5A, 0x5A), space_after=6)

    # promo
    add_p(doc, "Акция «Больше, чем кешбэк!»", size=12, bold=True, space_after=4)
    yellow_underline(doc.paragraphs[-1])
    promo = doc.add_table(rows=1, cols=1)
    borders(promo, "F0C14A")
    cell = promo.rows[0].cells[0]
    shade(cell, "1A1A1A")
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    r = p.add_run("5 часов линии консультаций — в подарок")
    run_font(r, size=11, bold=True, color=DocRGB(0xF0, 0xC1, 0x4A))
    for line in [
        f"При оплате пакета личных кабинетов мы начисляем 5 часов линии консультаций бесплатно (выгода {GIFT_VALUE:,} ₽).".replace(",", " "),
        "Из них в среднем 4 часа уходят на настройку и запуск сервиса.",
        "Ещё 1 час остаётся у вас на вопросы, консультации и решение возникающих задач по 1С с нашими специалистами.",
        "Отдельно за часы линии консультаций платить не нужно — они полностью покрываются подарком.",
    ]:
        p = cell.add_paragraph()
        r = p.add_run(line)
        run_font(r, size=9, color=DocRGB(0xFF, 0xFF, 0xFF))

    add_p(doc, "Ваш пакет: 750 личных кабинетов на 12 месяцев", size=12, bold=True, space_after=4)
    yellow_underline(doc.paragraphs[-1])
    add_p(doc, "Состав лицензий: 500 + 200 + 50 кабинетов. Цена за сотрудника — около 25 ₽ в месяц.", size=9, color=DocRGB(0x5A, 0x5A, 0x5A))

    price = doc.add_table(rows=1, cols=2)
    borders(price, "F0C14A")
    a, b = price.rows[0].cells
    shade(a, "FFF8E8")
    shade(b, "FFF8E8")
    a.paragraphs[0].clear()
    r = a.paragraphs[0].add_run("1С:Кабинет сотрудника — 750 кабинетов / 12 месяцев")
    run_font(r, size=10, bold=True)
    b.paragraphs[0].clear()
    b.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = b.paragraphs[0].add_run(f"{TOTAL:,} ₽".replace(",", " "))
    run_font(r, size=14, bold=True)
    add_p(doc, "Настройка, запуск и час поддержки по вопросам 1С — бесплатно в рамках подарочных часов.", size=9, color=DocRGB(0x5A, 0x5A, 0x5A), space_after=10)

    add_p(doc, "1. Условия", size=12, bold=True, space_after=4)
    yellow_underline(doc.paragraphs[-1])
    for t in [
        "Подключаем сервис «1С:Кабинет сотрудника» на 750 личных кабинетов сроком на 12 месяцев.",
        "В запуск входит: подключение к вашей 1С, настройка ролей и правил, помощь с выпуском электронных подписей, запуск пилотной группы и инструкции для сотрудников.",
        "5 часов линии консультаций начисляются в подарок. Около 4 часов — настройка и запуск; 1 час остаётся на дополнительные вопросы и помощь по 1С.",
        "За часы линии консультаций клиент не платит.",
        "Электронная подпись для сотрудников (усиленная неквалифицированная) — бесплатно.",
        "Предоставляем шаблоны документов для перехода: положение, уведомления, согласия.",
        "Работы выполняются удалённо. Выезд и нетиповые доработки — по отдельному согласованию.",
        "Коммерческое предложение действует 30 календарных дней с даты направления.",
    ]:
        add_check(doc, t)

    add_p(doc, "2. Сроки", size=12, bold=True, space_after=4)
    yellow_underline(doc.paragraphs[-1])
    add_p(doc, "Внедрение простое: без отдельной платформы и без долгого проектного цикла.", size=9, color=DocRGB(0x5A, 0x5A, 0x5A))
    t = doc.add_table(rows=1, cols=3)
    borders(t)
    for i, h in enumerate(["Этап", "Срок", "Результат"]):
        cell = t.rows[0].cells[i]
        shade(cell, "1A1A1A")
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        run_font(r, size=9, bold=True, color=DocRGB(0xFF, 0xFF, 0xFF))
    for a, b, d in [
        ("Договор и доступы к 1С", "2–4 рабочих дня", "Можно начинать настройку"),
        ("Настройка и запуск (около 4 часов)", "1 рабочий день", "Сервис подключен, роли и подписи настроены"),
        ("Пилот на небольшой группе", "3–7 рабочих дней", "Проверен рабочий сценарий"),
        ("Обучение ключевых сотрудников", "параллельно", "Видеоуроки и короткие инструкции"),
        ("Подключение всех 750 сотрудников", "2–3 недели", "Массовый запуск личных кабинетов"),
        ("Резерв 1 час поддержки", "по запросу", "Вопросы и помощь по 1С после запуска"),
    ]:
        row = t.add_row().cells
        for i, val in enumerate((a, b, d)):
            row[i].paragraphs[0].clear()
            r = row[i].paragraphs[0].add_run(val)
            run_font(r, size=8, bold=(i < 2))
    add_p(doc, "Ориентир: настройка за 1 день · пилот за 1–2 недели · все 750 кабинетов за 1–1,5 месяца.", size=9, bold=True, space_after=10)

    add_p(doc, "3. Бюджет", size=12, bold=True, space_after=4)
    yellow_underline(doc.paragraphs[-1])
    bt = doc.add_table(rows=1, cols=3)
    borders(bt)
    for i, h in enumerate(["Статья", "Пояснение", "Сумма"]):
        cell = bt.rows[0].cells[i]
        shade(cell, "1A1A1A")
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        run_font(r, size=9, bold=True, color=DocRGB(0xFF, 0xFF, 0xFF))
    for a, b, d, gift in [
        ("Личные кабинеты, 750 шт. / 12 месяцев", "Пакеты 500 + 200 + 50", f"{TOTAL:,} ₽".replace(",", " "), False),
        ("Настройка и запуск (~4 часа)", "Из подарочных часов линии консультаций", "0 ₽", True),
        ("Резерв на вопросы по 1С (1 час)", "Из подарочных часов линии консультаций", "0 ₽", True),
        ("Подарок: 5 часов линии консультаций", f"Акция «Больше, чем кешбэк!» · выгода {GIFT_VALUE:,} ₽".replace(",", " "), "0 ₽", True),
    ]:
        row = bt.add_row().cells
        for i, val in enumerate((a, b, d)):
            if gift:
                shade(row[i], "FFF8E8")
            row[i].paragraphs[0].clear()
            r = row[i].paragraphs[0].add_run(val)
            run_font(r, size=8, bold=(i == 0 or i == 2), color=DocRGB(0x1F, 0x6B, 0x2E) if (gift and i == 2) else DocRGB(0x1A, 0x1A, 0x1A))

    total_t = doc.add_table(rows=1, cols=2)
    borders(total_t, "F0C14A")
    a, b = total_t.rows[0].cells
    shade(a, "1A1A1A")
    shade(b, "1A1A1A")
    a.paragraphs[0].clear()
    r = a.paragraphs[0].add_run("Итого к оплате")
    run_font(r, size=11, bold=True, color=DocRGB(0xF0, 0xC1, 0x4A))
    p = a.add_paragraph()
    r = p.add_run("Только пакет кабинетов. Часы линии — бесплатно.")
    run_font(r, size=8, color=DocRGB(0xFF, 0xFF, 0xFF))
    b.paragraphs[0].clear()
    b.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = b.paragraphs[0].add_run(f"{TOTAL:,} ₽".replace(",", " "))
    run_font(r, size=16, bold=True, color=DocRGB(0xF0, 0xC1, 0x4A))

    add_p(doc, "Сравнение с другими решениями", size=12, bold=True, space_after=4)
    yellow_underline(doc.paragraphs[-1])
    ct = doc.add_table(rows=1, cols=3)
    borders(ct)
    for i, h in enumerate(["Критерий", "Форус + 1С:Кабинет сотрудника", "Другие операторы"]):
        cell = ct.rows[0].cells[i]
        shade(cell, "F0C14A" if i == 1 else "1A1A1A")
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        run_font(r, size=8, bold=True, color=DocRGB(0x1A, 0x1A, 0x1A) if i == 1 else DocRGB(0xFF, 0xFF, 0xFF))
    for a, b, d in [
        ("Сложность запуска", "Типовая настройка около 4 часов", "Часто долгое внедрение новой системы"),
        ("Где работают сотрудники", "В привычной 1С и личном кабинете", "Отдельная кадровая платформа"),
        ("Нагрузка на ИТ-отдел", "Без новых клиентских лицензий 1С", "Новые доступы, интеграции, сопровождение"),
        ("Где хранятся документы", "В вашей базе 1С", "Облако оператора, часто за отдельную плату"),
        ("Оплата запуска", "Часы поддержки — в подарок", "Внедрение обычно оплачивается отдельно"),
    ]:
        row = ct.add_row().cells
        shade(row[1], "EAF6EA")
        shade(row[2], "F9EAEA")
        for i, val in enumerate((a, b, d)):
            row[i].paragraphs[0].clear()
            r = row[i].paragraphs[0].add_run(val)
            col = DocRGB(0x1F, 0x6B, 0x2E) if i == 1 else (DocRGB(0x8C, 0x20, 0x20) if i == 2 else DocRGB(0x1A, 0x1A, 0x1A))
            run_font(r, size=8, bold=(i == 0), color=col)

    add_p(doc, "", space_after=6)
    add_p(doc, "Готовы подключить 750 кабинетов и помочь с запуском", size=11, bold=True, align="center")
    add_p(doc, "Проведём демонстрацию для кадровой службы и ИТ-отдела и ответим на вопросы.", size=9, align="center", color=DocRGB(0x5A, 0x5A, 0x5A))

    mt = doc.add_table(rows=1, cols=2)
    borders(mt, "F0C14A")
    a, b = mt.rows[0].cells
    shade(a, "FFF8E8")
    shade(b, "FFF8E8")
    a.paragraphs[0].clear()
    r = a.paragraphs[0].add_run("Ваш менеджер")
    run_font(r, size=8, color=DocRGB(0x5A, 0x5A, 0x5A))
    p = a.add_paragraph()
    r = p.add_run(MANAGER)
    run_font(r, size=12, bold=True)
    for line in [MANAGER_EMAIL, MANAGER_PHONE]:
        p = a.add_paragraph()
        r = p.add_run(line)
        run_font(r, size=9)
    b.paragraphs[0].clear()
    r = b.paragraphs[0].add_run("ГК «Форус»")
    run_font(r, size=11, bold=True)
    for line in ["www.forus.ru", "г. Иркутск, ул. Ямская, 1/1"]:
        p = b.add_paragraph()
        r = p.add_run(line)
        run_font(r, size=9, color=DocRGB(0x5A, 0x5A, 0x5A))

    doc.save(OUT_DOCX)
    print("DOCX:", OUT_DOCX)


def main():
    build_pdf()
    build_docx()
    # copies to root
    for src, name in [
        (OUT_PDF, "КП_ТФМ_Спецтехника_КЭДО_750.pdf"),
        (OUT_DOCX, "КП_ТФМ_Спецтехника_КЭДО_750.docx"),
        (OUT_PDF, "КП_Кабинет_сотрудника_750.pdf"),
    ]:
        dst = ROOT / name
        dst.write_bytes(src.read_bytes())
        print("copy", dst)


if __name__ == "__main__":
    main()
